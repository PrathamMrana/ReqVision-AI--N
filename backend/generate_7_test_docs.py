import os
from docx import Document

os.makedirs('tests/test_docs', exist_ok=True)

# 1. 01_BRD_Online_Library.docx
doc1 = Document()
doc1.add_heading('Business Requirements Document - Online Library Platform', 0)
doc1.add_paragraph('Business Objective: Deliver a comprehensive enterprise digital library solution.')
doc1.add_paragraph('Stakeholders: Library Director, Chief Technology Officer, Head of Member Services.')
doc1.add_paragraph('BR-001: The system shall provide a unified online catalog for members to search, browse, and discover digital books.')
doc1.add_paragraph('BR-002: The platform shall support automated late fee calculations and online payment settlement for overdue items.')
doc1.add_paragraph('BR-003: The system must enforce role-based access control for Librarians, Members, and System Administrators.')
doc1.add_paragraph('BR-004: The platform shall generate monthly inventory and circulation reports for library management analytics.')
doc1.add_paragraph('BR-005: The system shall provide secure user authentication and account credential management.')
doc1.add_paragraph('BR-006: The platform shall allow members to renew active book loans online before the due date.')
doc1.add_paragraph('BR-008: The system shall dispatch automated notification alerts to members regarding upcoming loan due dates.')
doc1.add_paragraph('BR-009: The platform shall maintain an immutable audit trail for all circulation status changes and financial transactions.')
doc1.add_paragraph('BR-010: The platform must support high traffic throughput during university examination periods.') # Intentionally unmapped
doc1.save('tests/test_docs/01_BRD_Online_Library.docx')

# 2. 02_SRS_Online_Library.docx
doc2 = Document()
doc2.add_heading('Software Requirements Specification - Online Library System', 0)
doc2.add_paragraph('Scope: System functional requirements and engineering specifications.')
doc2.add_paragraph('FR-101: The system shall provide a search module with metadata indexing to query books with sub-200ms response time.')
doc2.add_paragraph('FR-102: The system shall calculate overdue fines daily and process payment settlement through integrated payment gateway.')
doc2.add_paragraph('FR-103: The system shall enforce role-based authorization matrix for Member, Librarian, and Administrator permissions.')
doc2.add_paragraph('FR-104: The system shall compile and export monthly circulation and inventory statistics in PDF and CSV formats.')
doc2.add_paragraph('FR-105: Users shall authenticate with email and password using salted one-way hash verification.')
doc2.add_paragraph('FR-106: The system shall allow members to execute loan renewals if no reservation hold exists on the book.')
doc2.add_paragraph('FR-107: The notification service shall dispatch email alerts 48 hours prior to book loan due dates via SMTP.')
doc2.add_paragraph('FR-108: The checkout service shall validate maximum allowable loan quota per member and provide mobile browser access.')
doc2.add_paragraph('FR-109: The logging module shall record all book status transitions and checkout events to an immutable audit table.')
doc2.add_paragraph('FR-111: The system shall interface with legacy magnetic tape archival storage.') # Intentionally unmapped
doc2.save('tests/test_docs/02_SRS_Online_Library.docx')

# 3. 03_FRD_Online_Library.docx
doc3 = Document()
doc3.add_heading('Functional Requirements Document - Subsystem Architecture', 0)
doc3.add_paragraph('Functional Specification: Component-level process flows and interface behaviors.')
doc3.add_paragraph('FS-201: Search service indexes book title, author, ISBN, and genre with search filters and pagination.')
doc3.add_paragraph('FS-202: Payment integration handles Stripe and credit card gateway fee settlement with transaction logs.')
doc3.add_paragraph('FS-203: Authorization middleware validates JWT claims and role permissions on every API request.')
doc3.add_paragraph('FS-204: Report engine queries monthly circulation data and formats data tables into printable documents.')
doc3.add_paragraph('FS-205: Authentication controller verifies member login credentials against encrypted user database.')
doc3.add_paragraph('FS-206: Loan renewal endpoint validates due date eligibility and increments loan duration by 14 days.')
doc3.add_paragraph('FS-207: Scheduled background cron triggers due-date reminder emails via SMTP server.')
doc3.add_paragraph('FS-208: Mobile client layout and quota validation check rejects checkout request if member has reached active loan limit.')
doc3.add_paragraph('FS-209: Audit interceptor persists actor ID, timestamp, and action payload to PostgreSQL audit log.')
doc3.add_paragraph('FS-210: Scalability and load balancing specification for high throughput requests.')
doc3.add_paragraph('FS-211: The service shall store passwords using reversible encryption so administrators can recover the original password.') # INTENTIONAL CONFLICT with FR-105
doc3.save('tests/test_docs/03_FRD_Online_Library.docx')

# 4. 04_User_Stories_Online_Library.docx
doc4 = Document()
doc4.add_heading('User Stories - Sprint Product Backlog', 0)
doc4.add_paragraph('US-301: As a Member, I want to search books by title or author so that I can quickly find my reading material.')
doc4.add_paragraph('US-302: As a Member, I want to pay my late return fines online so that my borrowing privileges are restored.')
doc4.add_paragraph('US-303: As a Librarian, I want to access restricted inventory controls so that I can catalog new books.')
doc4.add_paragraph('US-304: As a Library Director, I want to download monthly circulation reports so that I can monitor library usage.')
doc4.add_paragraph('US-305: As a User, I want to securely log in with my email and password so that I can access my profile.')
doc4.add_paragraph('US-306: As a Member, I want to renew my borrowed book online so that I avoid overdue penalty fees.')
doc4.add_paragraph('US-307: As a Member, I want to receive an email notification before my book is due so that I return it on time.')
doc4.add_paragraph('US-308: As a Member, I want to access my library loan quota from mobile browser.')
doc4.add_paragraph('US-309: As an Administrator, I want to inspect system audit logs so that I can track book status modifications.')
doc4.add_paragraph('US-310: As an Administrator, I want to export legacy catalog tape archive format.') # Intentionally unmapped
doc4.add_paragraph('US-311: As a User, I want to sign in with social OAuth providers.')
doc4.save('tests/test_docs/04_User_Stories_Online_Library.docx')

# 5. 05_Test_Cases_Online_Library.docx
doc5 = Document()
doc5.add_heading('Test Case Specification - Verification & QA Suite', 0)
doc5.add_paragraph('TC-401: Test Scenario: Catalog Search Execution. Verify search query returns matching book records under 200ms latency.')
doc5.add_paragraph('TC-402: Test Scenario: Search Filtering. Verify filtering books by author and ISBN returns accurate subset.')
doc5.add_paragraph('TC-403: Test Scenario: Overdue Fee Payment. Verify payment processing succeeds and resets outstanding fine balance.')
doc5.add_paragraph('TC-404: Test Scenario: Role-Based Authorization. Verify Member account cannot access Librarian management endpoints.')
doc5.add_paragraph('TC-405: Test Scenario: Monthly Report Generation. Verify circulation PDF export contains total loan metrics.')
doc5.add_paragraph('TC-406: Test Scenario: User Authentication. Verify valid email and password returns JWT authentication token.')
doc5.add_paragraph('TC-407: Test Scenario: Book Loan Renewal. Verify loan renewal extends due date when no hold exists.')
doc5.add_paragraph('TC-408: Test Scenario: Due Date Email Dispatch. Verify reminder email is sent 48 hours prior to loan expiration.')
doc5.add_paragraph('TC-409: Test Scenario: Mobile Browser Quota Validation. Verify mobile browser view displays active loan count.')
doc5.add_paragraph('TC-410: Test Scenario: Audit Logging. Verify checkout event inserts corresponding record into audit log table.')
doc5.add_paragraph('TC-411: Test Scenario: Cafeteria Barcode Scanner. Verify food barcode scan.') # Intentionally unmapped
doc5.save('tests/test_docs/05_Test_Cases_Online_Library.docx')

# 6. 06_Change_Request_Online_Library.docx
doc6 = Document()
doc6.add_heading('Change Request - Engineering Change Collection', 0)
doc6.add_paragraph('CR-501: Reduce catalogue search query response time from 2 seconds to 1 second.')
doc6.add_paragraph('CR-502: Enforce Multi-Factor Authentication (MFA) via TOTP for all Librarian and Admin logins.')
doc6.add_paragraph('CR-503: Update payment gateway to support Apple Pay and Google Pay digital wallets.')
doc6.add_paragraph('CR-504: Add multi-format circulation report export to include JSON and XML formats.')
doc6.add_paragraph('CR-505: Add push notifications in addition to email for due date alerts.')
doc6.add_paragraph('CR-507: Replace office printers in administration office.') # Intentionally unmapped
doc6.save('tests/test_docs/06_Change_Request_Online_Library.docx')

# 7. 07_Meeting_Minutes_Online_Library.docx
doc7 = Document()
doc7.add_heading('Meeting Minutes - Architecture Review Board', 0)
doc7.add_paragraph('MOM-601: Architecture Review Board reviewed search indexing latency and approved CR-501.')
doc7.add_paragraph('MOM-602: Security committee confirmed MFA rollout CR-502 in Sprint 14.')
doc7.add_paragraph('MOM-603: General administrative lunch schedule update.') # Intentionally unmapped
doc7.save('tests/test_docs/07_Meeting_Minutes_Online_Library.docx')

print('All 7 test documents successfully generated in tests/test_docs/!')
