from docx import Document
import PyPDF2
from reportlab.pdfgen import canvas
import os

os.makedirs('tests/mock_docs', exist_ok=True)

# 1. TXT: BRD
with open('tests/mock_docs/random_a.txt', 'w') as f:
    f.write("BUSINESS OBJECTIVE: Increase sales. STAKEHOLDERS: Marketing. BUSINESS NEED: BR-001 Fast checkout.")

# 2. DOCX: SRS
doc = Document()
doc.add_heading('System Requirements Specification', 0)
doc.add_paragraph('System shall authenticate users.')
doc.add_paragraph('FR-001 Users must login with email.')
doc.save('tests/mock_docs/abc123.docx')

# 3. PDF: User Story
c = canvas.Canvas('tests/mock_docs/document_xyz.pdf')
c.drawString(100, 750, "As a user, I want to add items to my cart.")
c.drawString(100, 730, "So that I can checkout later.")
c.save()

print("Test files generated!")
