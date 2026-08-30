"""
Generate 7 InsureFlow Insurance Claim & Policy Platform documents (Blind Dataset).

Domain: Auto & Health insurance FNOL claim filing, damage estimate computation,
adjuster appraisal approval workflow, claim withdrawal/cancellation, accident photo upload,
duplicate invoice fraud detection, claims auditor RBAC, finance claim payout reconciliation.

Intentionally unmapped/adversarial artifacts:
  US-310: "Export paper insurance claim records to microfiche format for underground storage"
  TC-410: "Verify adjuster breakroom HEPA air purifier ionizer activates on schedule"
  CR-506: "Install electric vehicle charging stations in claims adjuster parking lot"
  DEC-603: "Team discussed AI computer vision damage appraisals but did not define algorithm accuracy thresholds"
  DEC-607: "Claims facility director decided to procure 25 ergonomic motorized standing desks"

Run: python generate_insureflow_docs.py
Output: tests/insureflow_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "insureflow_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — InsureFlow Claims Platform", [
        "BR-001: The InsureFlow platform shall provide an online portal where policyholders can submit First Notice of Loss (FNOL) insurance claims.",
        "BR-002: The system shall calculate automated vehicle damage repair cost estimates based on insurance policy coverage limits and deductible rules.",
        "BR-003: Licensed claims adjusters shall review, appraise, and approve or reject claim reimbursement requests submitted by policyholders.",
        "BR-004: Policyholders shall be able to withdraw or cancel a pending insurance claim before adjuster appraisal finalization.",
        "BR-005: The mobile application shall capture and upload digital photos of vehicle accident damage and repair body shop invoices.",
        "BR-006: The fraud detection engine shall compute perceptual invoice hashes to detect and block duplicate repair invoice claims.",
        "BR-007: Role-based authorization shall enforce strict separation between policyholders, field adjusters, and fraud investigation auditors.",
        "BR-008: Finance staff shall reconcile approved claim settlement payouts against bank disbursement files and ledger records.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_InsureFlow.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — InsureFlow Claims Platform", [
        "FR-101: The claim intake service shall allow policyholders to enter accident date, incident location, involved parties, and submit FNOL claims.",
        "FR-102: The estimation engine shall compute projected claim payout amounts by applying deductible schedules and insurance policy limits.",
        "FR-103: The adjuster workflow service shall route pending claims to designated claims adjusters for damage appraisal and approval.",
        "FR-104: The claim withdrawal module shall permit policyholders to withdraw pending claims and release insurance reserve holds.",
        "FR-105: The mobile upload service shall allow policyholders to photograph vehicle damage, perform invoice OCR, and attach photos to claims.",
        "FR-106: The fraud prevention module shall compute image hashes and merchant totals to detect duplicate invoice submissions across claims.",
        "FR-107: The access control engine shall enforce RBAC boundaries preventing policyholders from viewing other policyholders' claim files.",
        "FR-108: The financial settlement service shall match approved claim disbursement records against corporate bank payout transaction feeds.",
        "FR-109: Policyholder account passwords shall be stored using reversible DES encryption so that customer support agents can read passwords over phone.",
        "FR-110: The claim processing platform shall support a minimum of 6000 concurrent claim filings during catastrophic severe weather storm events.",
        "NFR-101: The insurance claim platform shall comply with NAIC insurance data privacy and state statutory insurance regulations.",
        "NFR-102: The online claim submission portal shall maintain 99.95% system uptime availability during disaster declaration periods.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_InsureFlow.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — InsureFlow Claims Platform", [
        "FS-201: Implement FNOL intake REST API POST /api/claims; validate active policy coverage and create record in claims_intake table.",
        "FS-202: Implement claim estimation engine querying Mitchell / CCC ONE auto body repair rate databases and policy deductible tables.",
        "FS-203: Implement claims adjuster appraisal workflow with interactive appraisal forms, payout calculation, and digital approval sign-off.",
        "FS-204: Implement claim cancellation endpoint POST /api/claims/{id}/withdraw; transition status to WITHDRAWN and close claim file.",
        "FS-205: Implement mobile photo upload API POST /api/claims/photos; invoke Tesseract OCR worker for body shop repair invoice parsing.",
        "FS-206: Implement duplicate invoice detection using dHash perceptual image hashing and exact repair shop invoice number index.",
        "FS-207: Implement RBAC JWT security interceptor validating role claims: Policyholder < ClaimsAdjuster < SpecialInvestigationAuditor.",
        "FS-208: Implement bank payout reconciliation worker parsing NACHA / ACH settlement feeds and matching claim disbursement reference numbers.",
        "FS-209: Policyholder passwords shall be secured using Argon2id salted cryptographic hashing. Plaintext or reversible storage is strictly forbidden.",
        "FS-210: Implement Kubernetes horizontal pod autoscaler and Redis cache clustering to sustain 6000 concurrent claim filings at p99 under 1s.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_InsureFlow.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — InsureFlow Claims Platform", [
        "US-301: As an insured driver, I want to file a First Notice of Loss claim online after an accident so that my insurance claim is registered.",
        "US-302: As a policyholder, I want to view an estimated claim payout calculation showing my deductible deduction before submitting.",
        "US-303: As a claims adjuster, I want to review submitted vehicle damage photos and approve claim repair estimates from my dashboard.",
        "US-304: As a driver who settled damages privately, I want to withdraw my pending insurance claim so that no claim record affects my policy.",
        "US-305: As a policyholder, I want to photograph body shop repair receipts and damage photos with my phone to attach to my claim.",
        "US-306: As a policyholder, I want to log into my insurance account using my email and secure password to view my claim status.",
        "US-307: As an insurance fraud investigator, I want an immutable audit log of all claim approvals and disbursement payouts for state audits.",
        "US-308: As an insurance financial clerk, I want an automated reconciliation report matching approved claim payouts against bank ACH files.",
        "US-309: As a policyholder filing during a hurricane storm, I want the insurance portal to process my claim without timeout crashes.",
        "US-310: As an insurance corporate archivist, I want to export paper claim dossiers to microfiche film for underground storage vaults.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_InsureFlow.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — InsureFlow Claims Platform", [
        "TC-401: Verify that a policyholder can submit an FNOL claim with accident date and location, and the record is created in claims_intake.",
        "TC-402: Verify that the estimation engine correctly applies the 500 dollar collision deductible and calculates projected insurance payout.",
        "TC-403: Verify that a claims adjuster can approve an appraised claim, status transitions to APPROVED, and payout authorization is generated.",
        "TC-404: Verify that a policyholder can withdraw a pending claim, status updates to WITHDRAWN, and claim reserve amount is released.",
        "TC-405: Verify that uploading a damage photo extracts repair shop invoice metadata and associates photos with the insurance claim.",
        "TC-406: Verify that uploading a duplicate repair invoice image flags the claim with a fraud warning and blocks duplicate settlement payout.",
        "TC-407: Verify that attempting to access another policyholder's claim without adjuster permissions returns 403 Forbidden.",
        "TC-408: Verify that user passwords are saved as salted Argon2id hashes and no reversible or plaintext passwords exist in credentials table.",
        "TC-409: Verify that the claim submission intake service sustains 6000 concurrent claim filings with p99 response time under 1 second.",
        "TC-410: Verify that the adjuster breakroom HEPA air purifier ionizer powers on when pressing the filter cycle button.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_InsureFlow.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — InsureFlow Claims Platform", [
        "CR-501: Enhance FR-103 to support instant SMS and WhatsApp claim approval notifications to policyholders in addition to email.",
        "CR-502: Modify FR-109 to prohibit reversible DES encryption and mandate salted Argon2id hashing in compliance with state regulations.",
        "CR-503: Extend FR-110 to support 12000 concurrent claim filings for the multi-state national flood catastrophe response.",
        "CR-504: Enhance FR-108 to support real-time Federal Reserve FedNow instant settlement bank reconciliation feeds.",
        "CR-505: Update FS-205 to support live video appraisal inspection streaming between policyholder and remote claims adjuster.",
        "CR-506: Install commercial electric vehicle Level 2 charging stations in the claims adjuster regional office parking lot.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_InsureFlow.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — InsureFlow Claims Platform", [
        "DEC-601: Committee agreed that FR-103 adjuster appraisal forms must support digital stylus customer signatures on tablet devices.",
        "DEC-602: Security lead decided that FR-109 password recovery practice must be revoked and replaced with self-service SMS token reset.",
        "DEC-603: The team discussed AI computer vision damage appraisals but did not define algorithm accuracy thresholds or validation rules.",
        "DEC-604: Finance director confirmed that FR-108 bank settlement reconciliation must run daily at 02:00 AM UTC against bank SFTP.",
        "DEC-605: Audit committee confirmed that FR-107 RBAC permission rules must undergo mandatory semi-annual compliance audits.",
        "DEC-606: QA team agreed that TC-409 storm surge load tests must be run on staging clusters replicating multi-region infrastructure.",
        "DEC-607: Claims facility director decided to procure 25 ergonomic motorized standing desks for regional field adjusters.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_InsureFlow.docx"))

    print(f"[InsureFlow] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
