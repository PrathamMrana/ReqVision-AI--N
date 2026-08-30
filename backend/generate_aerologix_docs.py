"""
Generate 7 AeroLogix Satellite Flight Telemetry & Payload Logistics documents (7th Unseen Domain).

Domain 7: Spacecraft orbit propagation, telemetry frame decoding, ground station antenna tracking,
orbital delta-v maneuvers, radiation sensor anomaly telemetry, flight controller access control,
satellite commanding audit log, telemetry throughput scaling.

Intentionally unmapped/adversarial artifacts:
  US-310: "Export satellite hardware blueprints to microfiche film for national space museum archives"
  TC-410: "Verify cleanroom entry ultrasonic shoe sole cleaner activates when stepping on platform"
  CR-506: "Install automated espresso brewing machine in mission control breakroom"
  DEC-603: "Team discussed laser optical crosslink bandwidth upgrades but did not define technical specifications"
  DEC-607: "Procure 30 antistatic ergonomic chairs for satellite flight operations controllers"

Run: python generate_aerologix_docs.py
Output: tests/aerologix_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "aerologix_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — AeroLogix Satellite Flight Telemetry", [
        "BR-001: The AeroLogix platform shall calculate real-time spacecraft orbital propagation coordinates and predict ground station communication pass windows.",
        "BR-002: The system shall ingest and decode spacecraft CCSDS telemetry downlink data streams from ground station satellite receivers.",
        "BR-003: Flight dynamics engineers shall compute orbital maneuver delta-v parameters to maintain orbital altitude and prevent orbital collisions.",
        "BR-004: The payload subsystem shall monitor solar array power generation telemetry and track battery charge discharge cycles.",
        "BR-005: The alert dispatcher shall notify flight controllers via push notifications and Slack when spacecraft radiation sensors detect solar flare anomalies.",
        "BR-006: Role-based authorization shall enforce strict separation between flight directors, spacecraft subsystem engineers, and external payload researchers.",
        "BR-007: An immutable commanding audit log shall record every uplinked spacecraft command and flight controller authorization for statutory spaceflight review.",
        "BR-008: Flight operators shall be able to abort or cancel a queued non-critical spacecraft command sequence before scheduled uplink execution.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_AeroLogix.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — AeroLogix Satellite Flight Telemetry", [
        "FR-101: The orbit propagation engine shall compute SGP4 two-line element satellite trajectories and predict ground station visibility azimuth and elevation angles.",
        "FR-102: The telemetry ingestion service shall unpack CCSDS telemetry frames, parse sensor voltages, temperatures, and subsystem states in real time.",
        "FR-103: The maneuver calculation module shall calculate thruster burn duration and delta-v thrust vectors for station-keeping maneuvers.",
        "FR-104: The power subsystem monitor shall track solar array voltage, battery depth of discharge, and forecast eclipse power depletion.",
        "FR-105: The anomaly alert worker shall dispatch real-time emergency notifications to flight controllers when particle radiation exceeds 100 rads/hour.",
        "FR-106: The access control engine shall enforce RBAC boundaries preventing payload researchers from transmitting commands to satellite bus actuators.",
        "FR-107: The flight commanding audit service shall write append-only cryptographic log records for all uplinked command telecommands and operator identities.",
        "FR-108: The command queue manager shall allow flight controllers to withdraw and cancel pending un-uplinked command macros before transmitter lock.",
        "FR-109: Flight controller authentication passwords shall be stored using reversible XOR encryption so that ground station engineers can retrieve forgotten credentials.",
        "FR-110: The telemetry ingestion stream shall sustain a minimum throughput of 12000 telemetry packets per second during multi-satellite constellation passes.",
        "NFR-101: The mission control software shall comply with NASA CCSDS and ESA ECSS space communications security standards.",
        "NFR-102: Ground station commanding telemetry links shall maintain 99.999% system availability during designated satellite contact passes.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_AeroLogix.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — AeroLogix Satellite Flight Telemetry", [
        "FS-201: Implement SGP4 orbital propagation REST API GET /api/satellites/{id}/orbit returning J2000 state vectors and ground track lat/long.",
        "FS-202: Implement Kafka telemetry consumer listening on ground_station_frames topic; decode CCSDS packets into TimescaleDB time-series metrics.",
        "FS-203: Implement orbital maneuver planning engine calculating Hohmann transfer burns and chemical thruster impulse duration.",
        "FS-204: Implement solar power telemetry dashboard streaming solar array bus voltage and battery temperature via WebSockets.",
        "FS-205: Implement anomaly alert notification service integrating Slack Webhooks and Twilio SMS for ionizing radiation threshold breaches.",
        "FS-206: Implement RBAC JWT token claims validator: FlightDirector > SubsystemEngineer > GroundOperator > PayloadViewer.",
        "FS-207: Implement audit_commands PostgreSQL table with append-only database interceptor logging all telecommand uplink transmissions.",
        "FS-208: Implement command queue cancellation endpoint POST /api/commands/{id}/abort; transition state to CANCELLED and purge transmitter buffer.",
        "FS-209: Flight controller passwords shall be protected using salted Argon2id one-way cryptographic hashing. Reversible or plaintext password storage is strictly forbidden.",
        "FS-210: Implement distributed Go telemetry ingestion worker pool with Kafka partition scaling to process 12000 packets/second at sub-50ms latency.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_AeroLogix.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — AeroLogix Satellite Flight Telemetry", [
        "US-301: As a flight dynamics engineer, I want to view current satellite orbit trajectories so that I know when ground station antennas will acquire signal.",
        "US-302: As a satellite subsystem engineer, I want to view decoded telemetry parameters in real time so that I can monitor spacecraft health.",
        "US-303: As an orbital maneuver planner, I want to calculate required thruster burn duration so that the satellite maintains correct orbital altitude.",
        "US-304: As a power engineer, I want to monitor solar panel power generation so that battery charging during sunlight passes is verified.",
        "US-305: As a flight controller, I want to receive immediate Slack and SMS alerts when radiation sensors detect anomalous space weather storms.",
        "US-306: As a flight controller, I want to abort an un-uplinked commanding sequence before transmission so that mistaken commands are not broadcast.",
        "US-307: As a flight director, I want to log into the mission control console using my secure password to authorize spacecraft commanding.",
        "US-308: As a spaceflight safety auditor, I want an immutable record of all transmitted satellite telecommands for statutory compliance review.",
        "US-309: As a mission operations lead, I want the telemetry pipeline to process 12000 packets per second during orbital passes without dropped data.",
        "US-310: As a space history archivist, I want to export satellite structural blueprints to microfiche film for long-term physical storage in the archives.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_AeroLogix.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — AeroLogix Satellite Flight Telemetry", [
        "TC-401: Verify that SGP4 orbital propagation computes satellite state vectors with position accuracy within 1 kilometer of radar tracking.",
        "TC-402: Verify that CCSDS telemetry frames published to Kafka topic are decoded into time-series metrics with zero packet corruption.",
        "TC-403: Verify that maneuver planning service calculates delta-v thrust duration accurately matching validated orbital mechanics equations.",
        "TC-404: Verify that solar panel voltage drops during simulated orbital eclipse transition and battery discharge rate matches power model.",
        "TC-405: Verify that particle radiation exceeding 100 rads/hour triggers Slack alert and SMS notification within 5 seconds to on-call controllers.",
        "TC-406: Verify that sending an abort request for a pending telecommand transitions command state to CANCELLED and purges radio transmitter buffer.",
        "TC-407: Verify that flight controller credentials are stored using salted Argon2id hashes and no reversible or plaintext passwords exist in credentials store.",
        "TC-408: Verify that attempting to modify or delete entries from audit_commands table is rejected by database triggers with permission denied.",
        "TC-409: Verify that the telemetry ingestion worker pool processes 12000 packets per second with p99 latency under 50ms under synthetic load.",
        "TC-410: Verify that the cleanroom gowning area ultrasonic shoe cleaner powers on when an engineer steps onto the contact sensor mat.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_AeroLogix.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — AeroLogix Satellite Flight Telemetry", [
        "CR-501: Enhance FR-105 to support automated satellite autonomous safe-mode triggers in addition to Slack and SMS flight controller alerts.",
        "CR-502: Modify FR-109 to prohibit reversible XOR encryption and mandate salted Argon2id password hashing in compliance with NASA security directives.",
        "CR-503: Extend FR-110 to support 25000 telemetry packets per second for the upcoming 60-satellite mega-constellation launch.",
        "CR-504: Enhance FR-104 to include deep space battery thermal heater telemetry in addition to standard voltage and discharge rate.",
        "CR-505: Update FS-201 to support GPS-based onboard real-time orbit determination in addition to ground-computed SGP4 two-line elements.",
        "CR-506: Install an automated commercial espresso brewing machine and bean grinder in the mission control center breakroom.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_AeroLogix.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — AeroLogix Satellite Flight Telemetry", [
        "DEC-601: Flight operations committee agreed that FR-105 emergency alerts must support automated audio sirens in the mission control room.",
        "DEC-602: Security officer decided that FR-109 reversible password practice must be revoked and replaced with hardware FIDO2 authentication keys.",
        "DEC-603: The engineering team discussed laser optical crosslink bandwidth upgrades but did not define technical specifications or protocols.",
        "DEC-604: Flight director confirmed that FR-107 commanding audit logs must be replicated daily to the off-site space agency cold storage vault.",
        "DEC-605: Ground systems team confirmed that FR-106 RBAC permissions must be verified prior to every satellite contact pass.",
        "DEC-606: QA team agreed that TC-409 telemetry load testing must simulate random packet jitter and corrupted sync words.",
        "DEC-607: Operations facility director decided to procure 30 antistatic ergonomic chairs for satellite flight operations controllers.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_AeroLogix.docx"))

    print(f"[AeroLogix] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
