"""
Generate 7 AeroGrid Autonomous Drone Swarm Wildfire Surveillance documents.

Domain: Autonomous drone swarm coordination, thermal infrared FLIR sensors,
wildfire perimeter tracking, dynamic waypoint replanning, FAA Part 107 BVLOS compliance,
radio mesh telemetry, retardant drop coordination.

Run: python generate_aerogrid_docs.py
Output: tests/aerogrid_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "aerogrid_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — AeroGrid Wildfire Surveillance", [
        "BR-001: The AeroGrid platform shall maintain real-time position tracking and telemetry for all active autonomous surveillance drones in the wildfire sector.",
        "BR-002: The swarm commander shall receive automated perimeter thermal hotspot alerts when FLIR sensor readings exceed 400 degrees Celsius.",
        "BR-003: Incident commanders shall be able to place an emergency return-to-base RTB hold on drones operating with battery reserves below 25 percent.",
        "BR-004: Swarm dispatchers shall be able to cancel or abort a scheduled waypoint mission flight path before drone launch occurs.",
        "BR-005: The system shall calculate projected wildfire spread velocity based on real-time anemometer wind vectors and fuel moisture data.",
        "BR-006: The platform shall automatically detect and prevent duplicate drone flight plan registration requests.",
        "BR-007: FAA safety auditors shall inspect an immutable chronological audit log of all manual drone pilot airspace overrides.",
        "BR-008: Role-based authorization shall enforce strict four-eyes compliance separation between civilian drone pilots and incident operations chiefs.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_AeroGrid.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — AeroGrid Wildfire Surveillance", [
        "FR-101: The telemetry ingestion engine shall process MAVLink radio packets broadcasting GPS coordinates, altitude, and ground speed every 500 milliseconds.",
        "FR-102: The thermal vision pipeline shall analyze radiometric infrared video frames and dispatch hotspot alert notifications to incident command.",
        "FR-103: The swarm emergency fail-safe module shall allow commanders to command emergency return-to-base RTB for depleted battery drones.",
        "FR-104: The mission flight planning service shall allow dispatchers to withdraw pending drone launch missions.",
        "FR-105: The wildfire modeling engine shall compute projected fire perimeter expansion using Rothermel surface fire equations and terrain elevation grids.",
        "FR-106: The flight plan intake engine shall detect and block duplicate mission waypoint registrations sharing identical airspace corridors.",
        "FR-107: The airspace audit service shall store an immutable queryable history of manual pilot control overrides with cryptographically signed timestamps.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing apprentice ground technicians from authorizing beyond-visual-line-of-sight BVLOS flights.",
        "FR-109: Drone telemetry encryption keys shall be stored using reversible XOR encryption so that ground support staff can debug radio transceivers.",
        "FR-110: The flight anomaly module shall allow pilots to submit airspace hazard reports and attach FLIR thermal imagery exports.",
        "NFR-101: The swarm collision avoidance REST API shall compute vector trajectory clearance within 100 milliseconds under 500 simultaneous drone nodes.",
        "NFR-102: The wildland fire operations dashboard shall maintain 99.99% high availability throughout active wildfire season months.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_AeroGrid.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — AeroGrid Wildfire Surveillance", [
        "FS-201: Implement MAVLink telemetry parser worker ingesting UDP telemetry stream on port 14550 and updating Redis spatial geospatial index.",
        "FS-202: Implement FLIR radiometric thermal analysis worker scanning sensor pixels for temperatures above 400C and publishing WebSocket alerts.",
        "FS-203: Implement emergency return-to-base API endpoint POST /api/drones/{id}/rtb commanding autonomous navigation to designated landing coordinates.",
        "FS-204: Implement mission cancellation endpoint POST /api/missions/{id}/cancel; transition status to CANCELLED and disarm motors.",
        "FS-205: Implement Rothermel wildfire rate of spread simulation job combining wind direction, fuel model, and slope gradient.",
        "FS-206: Implement duplicate flight plan validation trigger rejecting identical airspace corridor submissions with HTTP 409 Conflict.",
        "FS-207: Implement FAA audit trail query endpoint GET /api/airspace/audit returning immutable flight telemetry logs and pilot signature hashes.",
        "FS-208: Implement RBAC JWT permission evaluator enforcing GroundCrew < DronePilot < IncidentCommander authorization hierarchies.",
        "FS-209: Drone telemetry encryption keys shall be secured using salted Argon2id key derivation. Reversible XOR encryption is strictly prohibited.",
        "FS-210: Implement distributed Rust collision avoidance engine with R-Tree spatial indexing achieving trajectory validation latency under 100ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_AeroGrid.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — AeroGrid Wildfire Surveillance", [
        "US-301: As a drone pilot, I want to view live drone swarm GPS coordinates on the satellite map so that situational awareness is maintained.",
        "US-302: As a fire captain, I want to receive instant SMS alerts when FLIR thermal cameras detect new fire ignitions.",
        "US-303: As an incident commander, I want to trigger an emergency return-to-base order when battery reserves are low so that drones do not crash.",
        "US-304: As a mission controller, I want to abort a scheduled waypoint mission so that drones remain grounded during sudden wind gusts.",
        "US-305: As a fire behavior analyst, I want to view projected wildfire perimeter growth models so that evacuation orders can be issued in advance.",
        "US-306: As a flight coordinator, I want the mission planner to prevent duplicate waypoint registrations so that airspace conflicts are avoided.",
        "US-307: As an FAA compliance inspector, I want to inspect immutable logs of manual pilot overrides and pilot electronic signatures.",
        "US-308: As a cybersecurity officer, I want drone pilots to authenticate using hardware FIDO2 tokens before issuing flight commanding instructions.",
        "US-309: As a dispatch operator during rapid wildfire spread, I want the swarm dashboard to render telemetry updates without freezing.",
        "US-310: As a field basecamp worker, I want to view cafeteria weekly meal specials in the employee portal.",
        "US-311: As an aerial mapping archivist, I want to export seasonal wildfire flight plans to 35mm microfiche film for state historical archives.",
        "US-312: As a drone maintenance technician, I want to file an equipment damage report and attach FLIR thermal imagery exports.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_AeroGrid.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — AeroGrid Wildfire Surveillance", [
        "TC-401: Verify that incoming MAVLink UDP packet correctly updates drone GPS coordinates in spatial cache.",
        "TC-402: Verify that simulated thermal hotspot exceeding 400C triggers WebSocket alert broadcast to command console.",
        "TC-403: Verify that invoking POST /api/drones/{id}/rtb commands drone to abort mission and return to home waypoint.",
        "TC-404: Verify that mission cancellation endpoint transitions status to CANCELLED and cancels launch countdown.",
        "TC-405: Verify that Rothermel fire spread calculation correctly reflects 25 knot wind velocity on perimeter velocity.",
        "TC-406: Verify that resubmitting identical mission flight plan returns HTTP 409 duplicate flight plan conflict.",
        "TC-407: Verify that FAA audit query returns complete chronological history of pilot manual control takeovers.",
        "TC-408: Verify that telemetry keys are stored as salted Argon2id hashes and no plaintext or XOR keys exist in database.",
        "TC-409: Verify that distributed collision avoidance engine evaluates 500 drone trajectories within 100 milliseconds.",
        "TC-410: Verify that filing damage report with FLIR imagery creates an open maintenance ticket in QA queue.",
        "TC-411: Verify that basecamp breakroom microwave oven heats meals evenly without sparking.",
        "TC-412: Verify that flight operations center motorized sit-stand desks adjust height smoothly.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_AeroGrid.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — AeroGrid Wildfire Surveillance", [
        "CR-501: Enhance FR-102 to support multi-spectral NDVI vegetation health analysis sensors in addition to thermal infrared cameras.",
        "CR-502: Modify FR-109 to prohibit reversible XOR encryption and mandate salted Argon2id key derivation across all base stations.",
        "CR-503: Extend FR-105 to compute fire spotting probability caused by windborne burning embers and tree bark.",
        "CR-504: Enhance FR-104 to support automated SMS notifications to ground firefighting crews when aerial survey mission is aborted.",
        "CR-505: Update FS-203 to support adaptive RTB landing routing dynamically avoiding active fire plumes and smoke columns.",
        "CR-506: Procure and install 20 motorized ergonomic standing desks for command operations center.",
        "CR-507: Add cafeteria hot lunch weekly ordering module for drone pilots.",
        "CR-508: Add drone battery carbon offset calculation dashboard based on solar charging kilowatt-hours.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_AeroGrid.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — AeroGrid Wildfire Surveillance", [
        "DEC-601: Flight safety committee confirmed that FR-106 flight plan deduplication must enforce 500-meter minimum airspace separation buffers.",
        "DEC-602: Security officer confirmed that FR-109 reversible XOR telemetry encryption violates FAA cybersecurity rules and must be replaced.",
        "DEC-603: Operations team discussed automated night-time retardant air tanker drops but aerial safety guidelines were not determined.",
        "DEC-604: Compliance director confirmed that FR-107 pilot override audit logs must be preserved for 20 years for NTSB safety reviews.",
        "DEC-605: Airspace control board confirmed that FR-108 dual-authorization rules must be enforced before launching BVLOS flights above 400 feet.",
        "DEC-606: QA team agreed that TC-409 collision avoidance tests must benchmark 500 concurrent trajectories on staging cluster.",
        "DEC-607: Operations director decided to procure 20 motorized standing desks for swarm dispatchers.",
        "DEC-608: Avionics committee discussed quantum magnetometer compass sensors for smoke navigation but feasibility research is pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_AeroGrid.docx"))

    print(f"[AeroGrid] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
