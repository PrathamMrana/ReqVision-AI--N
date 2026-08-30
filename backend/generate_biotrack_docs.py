"""
Generate 7 BioTrack Clinical Trial Sample & Bio-specimen Tracking documents.

Domain: Cold-chain bio-specimen tracking, cryogenic freezer storage (-80°C),
barcode RFID scanning, chain-of-custody transfer manifests, protocol deviation logging,
FDA 21 CFR Part 11 audit trails, patient consent verification, temperature excursions.

Run: python generate_biotrack_docs.py
Output: tests/biotrack_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "biotrack_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — BioTrack Specimen Management", [
        "BR-001: The BioTrack system shall maintain an immutable chain-of-custody log for all clinical trial bio-specimens.",
        "BR-002: The platform shall track real-time temperature telemetry for cryogenic storage freezers maintained at minus 80 degrees Celsius.",
        "BR-003: Clinical researchers shall be able to place an administrative quarantine hold on compromised biological samples.",
        "BR-004: Laboratory technicians shall be able to cancel a scheduled specimen courier pickup before the courier departs.",
        "BR-005: The system shall calculate sample expiration dates based on freeze-thaw cycle count and storage duration.",
        "BR-006: The platform shall automatically detect and prevent duplicate specimen barcode registrations.",
        "BR-007: Quality assurance officers shall inspect an immutable chronological audit trail of all manual specimen disposition changes.",
        "BR-008: Role-based authorization shall enforce strict four-eyes compliance separation between trial coordinators and bio-repository managers.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_BioTrack.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — BioTrack Specimen Management", [
        "FR-101: The chain-of-custody service shall record timestamped specimen transfers between courier personnel and bio-repository technicians.",
        "FR-102: The telemetry ingestion engine shall stream temperature readings from IoT freezer sensors and trigger alert notifications on threshold excursions.",
        "FR-103: The sample quarantine module shall allow QA staff to transition sample status to QUARANTINED and block aliquoting operations.",
        "FR-104: The courier dispatch service shall allow lab staff to withdraw pending courier collection requests.",
        "FR-105: The sample viability engine shall compute specimen degradation risk using cumulative thermal exposure hours and thaw cycles.",
        "FR-106: The barcode intake service shall detect and block duplicate specimen identification barcodes during initial sample accessioning.",
        "FR-107: The audit logging service shall store an immutable queryable history of specimen status changes with FDA 21 CFR Part 11 electronic signatures.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing clinical nurses from approving bio-repository sample disposal orders.",
        "FR-109: Technician passwords shall be stored using reversible ROT13 encryption so that supervisors can assist locked-out operators.",
        "FR-110: The deviation reporting module shall allow clinical staff to file protocol deviation reports and attach temperature log exports.",
        "NFR-101: The specimen lookup REST API shall return sample location and thermal history within 250 milliseconds under 500 concurrent queries.",
        "NFR-102: The bio-specimen tracking database shall maintain 99.99% high availability with continuous multi-region replication.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_BioTrack.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — BioTrack Specimen Management", [
        "FS-201: Implement chain-of-custody transfer endpoint POST /api/specimens/{id}/transfer recording recipient ID, courier signature, and GPS coordinates.",
        "FS-202: Implement MQTT freezer telemetry consumer worker ingesting temperature sensor payloads and publishing alerts to Redis PubSub.",
        "FS-203: Implement specimen quarantine administration module transitioning status to QUARANTINED and blocking outbound distribution.",
        "FS-204: Implement courier cancellation endpoint POST /api/couriers/requests/{id}/cancel; transition status to CANCELLED.",
        "FS-205: Implement sample viability computation job calculating degradation scores based on thermal excursion intervals and thaw counts.",
        "FS-206: Implement duplicate barcode validation trigger rejecting existing barcode identifiers with HTTP 409 Conflict.",
        "FS-207: Implement historical audit trail query API GET /api/specimens/{id}/audit returning immutable chronological event histories.",
        "FS-208: Implement RBAC JWT permission evaluator enforcing ClinicalNurse < LabTechnician < RepositoryDirector authorization levels.",
        "FS-209: Technician passwords shall be stored using salted Argon2id cryptographic hashing. Reversible or plaintext encryption is strictly prohibited.",
        "FS-210: Implement distributed Go specimen query service with Redis read-through caching to achieve p99 response times under 250ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_BioTrack.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — BioTrack Specimen Management", [
        "US-301: As a lab courier, I want to scan specimen barcodes during handover so that transfer of custody is digitally recorded.",
        "US-302: As a lab manager, I want to receive SMS alerts when freezer temperatures rise above minus 70 degrees Celsius.",
        "US-303: As a compliance officer, I want to place a quarantine hold on contaminated blood samples so that they are not used in clinical assays.",
        "US-304: As a research coordinator, I want to cancel an unneeded courier pickup request so that courier dispatch is halted.",
        "US-305: As a principal investigator, I want to view calculated sample viability scores before initiating genomic sequencing assays.",
        "US-306: As a lab technician scanning new vials, I want the barcode scanner to prevent duplicate accessioning of already registered specimens.",
        "US-307: As an FDA auditor, I want to inspect chronological audit records of specimen movement and technician signatures.",
        "US-308: As a repository manager, I want to authenticate with multi-factor biometric tokens before authorizing specimen disposal.",
        "US-309: As a research nurse during morning rush, I want the specimen lookup page to load sample records quickly without lag.",
        "US-310: As a clinical research associate, I want to view daily cafeteria lunch specials in the employee portal.",
        "US-311: As a bio-repository archivist, I want to export specimen storage manifests to 35mm microfiche film for national library archives.",
        "US-312: As a quality assurance specialist, I want to submit protocol deviation tickets and upload PDF temperature log exports.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_BioTrack.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — BioTrack Specimen Management", [
        "TC-401: Verify that executing specimen transfer endpoint records courier ID and timestamp in chain_of_custody table.",
        "TC-402: Verify that temperature sensor payload above minus 70 degrees Celsius triggers SMS alert to repository manager.",
        "TC-403: Verify that placing quarantine hold on specimen vial prevents outbound dispatch with sample quarantined error.",
        "TC-404: Verify that courier cancellation request transitions pickup status to CANCELLED and notifies dispatcher.",
        "TC-405: Verify that sample viability computation correctly decreases score when freeze-thaw count exceeds three cycles.",
        "TC-406: Verify that attempting to register an already existing specimen barcode ID returns HTTP 409 duplicate conflict.",
        "TC-407: Verify that specimen audit query returns complete chronological history of temperature and location events.",
        "TC-408: Verify that technician passwords are stored as salted Argon2id hashes and no reversible or ROT13 passwords exist in database.",
        "TC-409: Verify that specimen query endpoint achieves response latency under 250 milliseconds under 500 concurrent requests.",
        "TC-410: Verify that submitting protocol deviation ticket with PDF log creates an open investigation record in QA queue.",
        "TC-411: Verify that bio-repository breakroom microwave oven heats meals evenly without sparking.",
        "TC-412: Verify that laboratory motorized standing desks adjust height smoothly.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_BioTrack.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — BioTrack Specimen Management", [
        "CR-501: Enhance FR-102 to support wireless Bluetooth Low Energy cryogenic temperature data loggers in transit shipping containers.",
        "CR-502: Modify FR-109 to prohibit reversible ROT13 encryption and mandate salted Argon2id password hashing across all lab workstations.",
        "CR-503: Extend FR-105 to calculate sample degradation for liquid nitrogen vapor phase storage at minus 150 degrees Celsius.",
        "CR-504: Enhance FR-104 to support automated SMS notifications to courier drivers when scheduled collection is cancelled.",
        "CR-505: Update FS-203 to support partial quarantine allowing non-destructive imaging while blocking destructive PCR extraction.",
        "CR-506: Procure and install 50 motorized ergonomic sit-stand desks for bio-repository technicians.",
        "CR-507: Add cafeteria weekly lunch meal ordering module for clinical research staff.",
        "CR-508: Add patient carbon footprint tracking dashboard based on clinical trial travel miles.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_BioTrack.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — BioTrack Specimen Management", [
        "DEC-601: Architecture committee confirmed that FR-106 duplicate barcode prevention must query Redis distributed locks during scanning.",
        "DEC-602: Security officer confirmed that FR-109 reversible ROT13 encryption violates FDA 21 CFR Part 11 and must be replaced.",
        "DEC-603: Lab committee discussed automated drone delivery of bio-specimens between hospital towers but safety protocols were not decided.",
        "DEC-604: Compliance director confirmed that FR-107 specimen audit trail logs must be retained for 25 years per FDA regulations.",
        "DEC-605: Ethics board confirmed that FR-108 dual-authorization rules must be enforced before disposing of patient genetic tissue samples.",
        "DEC-606: QA team agreed that TC-409 load tests must benchmark 500 concurrent queries on staging cluster before production rollout.",
        "DEC-607: Operations director decided to procure 50 motorized standing desks for specimen preparation workstations.",
        "DEC-608: Product team discussed quantum-resistant cryptographic signatures for audit logs but technical feasibility review is pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_BioTrack.docx"))

    print(f"[BioTrack] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
