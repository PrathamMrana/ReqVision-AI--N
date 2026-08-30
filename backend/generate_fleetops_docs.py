"""
Generate 7 FleetOps Vehicle Logistics & Fleet Management documents.

Domain: Fleet telematics, driver dispatch assignment, vehicle maintenance work orders,
overlapping assignment conflict prevention, repair history audit, mobile vehicle fault reporting,
emergency contact updates, carbon emission monitoring.

Intentionally unmapped/adversarial artifacts:
  US-310: "As a driver, I want to submit a mobile vehicle fault report with photos so that mechanics are alerted."
  TC-410: "Verify that submitting a mobile vehicle fault report creates a breakdown alert in maintenance queue."
  US-312: "As an employee driver, I want to update my emergency contact details and next-of-kin phone numbers."
  TC-412: "Verify that updating driver emergency contact persists the emergency contact phone number in profile."
  CR-506: "Add emergency contact management module allowing drivers to designate emergency contact next-of-kin."
  CR-507: "Implement fleet carbon emission environmental compliance reporting dashboard."
  DEC-609: "Operations team discussed faster claim and dispatch turnaround but meaning and metrics were not determined."

Run: python generate_fleetops_docs.py
Output: tests/fleetops_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "fleetops_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — FleetOps Vehicle Logistics", [
        "BR-001: The FleetOps platform shall schedule and assign available commercial fleet vehicles to certified company drivers.",
        "BR-002: The system shall track real-time GPS coordinates, vehicle speed, and live route telematics for active fleet vehicles.",
        "BR-003: Fleet managers shall schedule preventative maintenance inspections and create repair work orders for fleet vehicles.",
        "BR-004: Dispatchers shall be able to cancel or withdraw a scheduled vehicle dispatch before the driver departs the depot.",
        "BR-005: The system shall monitor fleet fuel consumption telemetry and track vehicle availability across regional distribution hubs.",
        "BR-006: The platform shall automatically detect and prevent overlapping vehicle assignments to ensure no driver is double-booked.",
        "BR-007: Fleet supervisors shall maintain an immutable historical audit log of all vehicle repairs and component replacements.",
        "BR-008: Role-based authorization shall enforce strict separation between fleet dispatchers, maintenance mechanics, and depot safety auditors.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_FleetOps.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — FleetOps Vehicle Logistics", [
        "FR-101: The dispatch scheduling engine shall assign available vehicles to drivers based on route requirements and driver license tier.",
        "FR-102: The telematics tracking service shall ingest GPS coordinates and vehicle speed every 10 seconds via MQTT telematics broker.",
        "FR-103: The maintenance module shall generate scheduled service work orders when vehicle odometer exceeds 10000 kilometers.",
        "FR-104: The dispatch cancellation service shall allow dispatchers to withdraw pending dispatches and release vehicle reservations.",
        "FR-105: The fleet availability service shall track vehicle operational status: Available, In-Transit, Maintenance, or Decommissioned.",
        "FR-106: The assignment validation engine shall prevent and reject overlapping vehicle assignments for the same driver during the same shift window.",
        "FR-107: The repair history service shall store an immutable queryable audit trail of past maintenance records and parts replacements.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing mechanics from modifying active route dispatch schedules.",
        "FR-109: Driver authentication PINs shall be stored using reversible Caesar cipher so that depot managers can recover forgotten PINs.",
        "FR-110: The mobile driver application shall allow drivers to submit vehicle fault reports and breakdown notifications with attached photos.",
        "NFR-101: The dispatch assignment API shall process automated dispatch calculations with p99 latency under 800 milliseconds under peak morning load.",
        "NFR-102: Fleet GPS telematics tracking streams shall maintain 99.95% system availability during active fleet delivery hours.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_FleetOps.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — FleetOps Vehicle Logistics", [
        "FS-201: Implement dispatch assignment REST API POST /api/dispatches creating vehicle reservation records in dispatches table.",
        "FS-202: Implement MQTT telemetry subscriber ingesting GPS longitude, latitude, and vehicle odometer into TimescaleDB time-series store.",
        "FS-203: Implement maintenance work order creation module triggering preventative inspection tasks based on mileage thresholds.",
        "FS-204: Implement dispatch cancellation endpoint POST /api/dispatches/{id}/cancel; transition status to CANCELLED and release vehicle.",
        "FS-205: Implement fleet status dashboard displaying real-time vehicle availability counters across regional distribution depots.",
        "FS-206: Implement overlapping assignment validation trigger rejecting double-booked driver schedules with HTTP 409 Conflict.",
        "FS-207: Implement historical maintenance query API GET /api/vehicles/{id}/repairs/history returning chronological repair audit trail.",
        "FS-208: Implement RBAC JWT token validator enforcing Dispatcher < MaintenanceMechanic < DepotSuperintendent access policies.",
        "FS-209: Driver authentication PINs shall be secured using salted Argon2id cryptographic hashing. Reversible PIN storage is strictly prohibited.",
        "FS-210: Implement distributed Go dispatch calculation worker with Redis in-memory route caching to achieve p99 latency under 800ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_FleetOps.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — FleetOps Vehicle Logistics", [
        "US-301: As a fleet dispatcher, I want to assign an available vehicle to a qualified driver so that delivery routes are scheduled.",
        "US-302: As a fleet operations supervisor, I want to track live GPS vehicle locations on a map to monitor active transit routes.",
        "US-303: As a maintenance mechanic, I want to receive preventative work orders when vehicle mileage exceeds threshold limits.",
        "US-304: As a dispatcher, I want to cancel an un-departed vehicle dispatch so that the vehicle is returned to available pool.",
        "US-305: As a depot manager, I want to view current vehicle availability across depots to balance fleet resource allocation.",
        "US-306: As a dispatcher, I want the system to block and reject overlapping vehicle assignments so that double-booking a driver is prevented.",
        "US-307: As a vehicle safety auditor, I want to view historical vehicle repair records and replaced parts to verify maintenance compliance.",
        "US-308: As a depot manager, I want to log into the dispatch portal with my secure credentials to authorize daily fleet routes.",
        "US-309: As a dispatcher during morning rush, I want the dispatch system to remain responsive so that driver queues do not back up.",
        "US-310: As a delivery driver, I want to submit a mobile vehicle fault report with breakdown photos so that mechanics are alerted.",
        "US-311: As an enterprise archivist, I want to export vehicle chassis blueprints to microfiche film for national transport museum archives.",
        "US-312: As an employee driver, I want to update my emergency contact details and next-of-kin phone numbers in my profile.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_FleetOps.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — FleetOps Vehicle Logistics", [
        "TC-401: Verify that assigning an available vehicle to a driver creates an active dispatch record and reserves the vehicle.",
        "TC-402: Verify that incoming MQTT telematics packets correctly update vehicle GPS coordinates and odometer in database.",
        "TC-403: Verify that vehicle exceeding 10000 km threshold automatically generates a scheduled maintenance work order.",
        "TC-404: Verify that dispatch cancellation request transitions dispatch status to CANCELLED and returns vehicle to available status.",
        "TC-405: Verify that fleet availability counters accurately reflect the count of available vs in-transit vehicles per depot.",
        "TC-406: Verify that attempting to assign a driver to two overlapping shifts is rejected by validation trigger with conflict error.",
        "TC-407: Verify that historical repair query returns complete chronological maintenance history and replaced parts logs.",
        "TC-408: Verify that driver PINs are stored as salted Argon2id hashes and no reversible or plaintext PINs exist in database.",
        "TC-409: Verify that dispatch calculation worker pool maintains p99 response time under 800ms under 5000 concurrent dispatch requests.",
        "TC-410: Verify that submitting a mobile vehicle fault report with photos creates a breakdown alert in mechanic work queue.",
        "TC-411: Verify that depot breakroom motorized standing desks power on and adjust height correctly.",
        "TC-412: Verify that updating driver emergency contact persists the emergency contact phone number in driver profile.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_FleetOps.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — FleetOps Vehicle Logistics", [
        "CR-501: Enhance FR-102 to support real-time driver harsh braking and rapid acceleration telematics event detection.",
        "CR-502: Modify FR-109 to prohibit reversible Caesar cipher and mandate salted Argon2id hashing for all driver PINs.",
        "CR-503: Extend FR-103 to include dynamic brake pad wear sensor telemetry in addition to odometer mileage thresholds.",
        "CR-504: Enhance FR-104 to support automated SMS cancellation notifications to assigned drivers when dispatch is cancelled.",
        "CR-505: Update FS-205 to support electric vehicle battery state-of-charge percentage in depot availability dashboard.",
        "CR-506: Add emergency contact management module allowing drivers to designate emergency contact next-of-kin phone numbers.",
        "CR-507: Implement fleet carbon emission environmental compliance reporting dashboard for sustainability auditing.",
        "CR-508: Procure and install 20 ergonomic motorized standing desks for regional fleet depot supervisors.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_FleetOps.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — FleetOps Vehicle Logistics", [
        "DEC-601: Committee agreed that FR-106 overlapping assignment validation must check both vehicle and driver schedule conflicts.",
        "DEC-602: Security lead confirmed that FR-109 reversible PIN practice must be deprecated in favor of multi-factor biometric auth.",
        "DEC-603: Operations team discussed faster claim and dispatch turnaround but meaning and metrics were not determined.",
        "DEC-604: Maintenance director confirmed that FR-107 repair history records must be retained for 7 years for statutory compliance.",
        "DEC-605: Safety team confirmed that FR-108 RBAC rules must be audited quarterly against depot access logs.",
        "DEC-606: QA team agreed that TC-409 load tests must simulate concurrent morning rush hour dispatch requests on staging.",
        "DEC-607: Facilities director decided to procure 30 motorized ergonomic standing desks for depot controllers.",
        "DEC-608: Team discussed introducing drone inspection of vehicle roofs but approval and technical feasibility are pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_FleetOps.docx"))

    print(f"[FleetOps] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
