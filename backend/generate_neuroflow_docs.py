"""
Generate 7 NeuroFlow Brain-Computer Interface (BCI) Telemetry & Neuromodulation documents.

Domain: Neural electroencephalography (EEG) signal processing, deep brain stimulation (DBS)
telemetry, real-time seizure prediction, adaptive closed-loop neuromodulation, FDA Class III medical device compliance,
artifact noise rejection, Bluetooth Low Energy (BLE) neural telemetry streaming.

Run: python generate_neuroflow_docs.py
Output: tests/neuroflow_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "neuroflow_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — NeuroFlow BCI Telemetry", [
        "BR-001: The NeuroFlow system shall continuously stream multichannel EEG sensor telemetry from wearable neural headbands to clinician monitoring dashboards.",
        "BR-002: The platform shall automatically detect pre-ictal neural spike patterns and dispatch early seizure alerts to caregivers.",
        "BR-003: Clinical neurologists shall be able to place an emergency stimulation pause on active deep brain stimulation neuro-stimulators.",
        "BR-004: Patients shall be able to cancel or abort a scheduled therapeutic electrical pulse stimulation session before electrode pulse delivery begins.",
        "BR-005: The system shall calculate real-time cortical power spectral density across delta, theta, alpha, and beta frequency bands.",
        "BR-006: The platform shall automatically detect and prevent duplicate neural telemetry packet transmissions from implanted bio-sensors.",
        "BR-007: Clinical compliance auditors shall inspect an immutable chronological audit trail of all manual electrode voltage calibration adjustments.",
        "BR-008: Role-based authorization shall enforce strict four-eyes medical compliance separation between clinical research assistants and certified neurosurgeons.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_NeuroFlow.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — NeuroFlow BCI Telemetry", [
        "FR-101: The neural telemetry ingestion engine shall process 256-channel EEG data streams sampled at 1000 Hz over encrypted Bluetooth Low Energy BLE links.",
        "FR-102: The seizure prediction classifier shall analyze spike-wave discharge synchrony and dispatch instant push notifications when pre-ictal probability exceeds 0.85.",
        "FR-103: The stimulation safety fail-safe module shall allow neurologists to command immediate stimulation halt and discharge electrode capacitance.",
        "FR-104: The therapy cancellation service shall allow patients to withdraw pending neuromodulation pulse delivery sequences.",
        "FR-105: The spectral analysis engine shall compute Fast Fourier Transform FFT power spectra across 0.5 to 30 Hz neural frequency bands.",
        "FR-106: The telemetry intake engine shall detect and block duplicate neural telemetry packets sharing identical sequence timestamp signatures.",
        "FR-107: The calibration audit service shall store an immutable queryable history of electrode voltage adjustments with cryptographic clinician signatures.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing non-certified lab technicians from modifying therapeutic stimulation parameters.",
        "FR-109: Patient neural biometric authentication templates shall be stored using reversible DES encryption so that support staff can recover forgotten patient profiles.",
        "FR-110: The clinical anomaly module shall allow neurologists to submit patient adverse event reports and attach EDF electroencephalogram recording exports.",
        "NFR-101: The real-time closed-loop neuromodulation API shall compute stimulation trigger decisions within 20 milliseconds under 100 concurrent patient streams.",
        "NFR-102: The clinical neurology platform shall maintain 99.999% high availability during active therapeutic neuromodulation operations.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_NeuroFlow.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — NeuroFlow BCI Telemetry", [
        "FS-201: Implement BLE telemetry ingestion service parsing 256-channel float32 microvolt arrays and buffering in RingBuffer queues.",
        "FS-202: Implement seizure prediction neural network inference pipeline broadcasting WebSocket alerts when pre-ictal index crosses 0.85.",
        "FS-203: Implement emergency stimulation stop endpoint POST /api/implants/{id}/emergency-stop disabling stimulation pulses in under 5ms.",
        "FS-204: Implement therapy cancellation endpoint POST /api/therapy/sessions/{id}/cancel; transition status to CANCELLED.",
        "FS-205: Implement FFT spectral density calculation batch worker computing band powers for delta, theta, alpha, and beta oscillations.",
        "FS-206: Implement duplicate packet rejection trigger discarding redundant neural packets with HTTP 409 Conflict status.",
        "FS-207: Implement voltage audit query endpoint GET /api/audit/calibration returning immutable chronological logs and clinician cryptographic signatures.",
        "FS-208: Implement RBAC JWT permission validator enforcing LabTech < ClinicalFellow < AttendingNeurosurgeon authorization tiers.",
        "FS-209: Patient biometric authentication templates shall be secured using salted Argon2id cryptographic hashing. Reversible DES encryption is strictly prohibited.",
        "FS-210: Implement C++ real-time closed-loop stimulation engine with SIMD vectorization achieving trigger latency under 20ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_NeuroFlow.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — NeuroFlow BCI Telemetry", [
        "US-301: As a clinical neurophysiologist, I want to observe live multichannel EEG waveforms on my workstation monitor so that brainwave abnormalities are visible.",
        "US-302: As a caregiver, I want to receive urgent smartphone alerts when pre-seizure neural indicators occur so that preventative medicine can be administered.",
        "US-303: As an attending neurologist, I want to execute an emergency stimulation shutdown when a patient reports headache aura so that adverse events are mitigated.",
        "US-304: As an epilepsy patient, I want to cancel an unneeded neuromodulation therapy session so that electrical pulses are not delivered.",
        "US-305: As a neural data researcher, I want to view calculated brainwave power spectral density graphs so that sleep oscillation stages can be categorized.",
        "US-306: As a biomedical engineer, I want the telemetry gateway to discard duplicate radio packet transmissions so that telemetry buffer overflow is prevented.",
        "US-307: As an FDA medical device auditor, I want to inspect immutable logs of electrode calibration changes and neurosurgeon electronic signatures.",
        "US-308: As a hospital security officer, I want neurologists to authenticate using hardware smartcards before accessing deep brain stimulator controls.",
        "US-309: As a clinical technician during high telemetry load, I want the EEG monitoring interface to render real-time waveforms smoothly without frame drops.",
        "US-310: As a neurology clinic receptionist, I want to view cafeteria weekly hot lunch specials in the staff portal.",
        "US-311: As a medical records archivist, I want to export historical patient EEG strip charts to 35mm microfiche film for state hospital archives.",
        "US-312: As a research neurologist, I want to submit a clinical adverse event report and attach EDF electroencephalogram recording exports.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_NeuroFlow.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — NeuroFlow BCI Telemetry", [
        "TC-401: Verify that 256-channel BLE telemetry stream correctly updates real-time multichannel EEG waveform display without data loss.",
        "TC-402: Verify that simulated pre-ictal spike activity exceeding 0.85 threshold triggers instant alert notification to caregiver mobile application.",
        "TC-403: Verify that emergency stimulation shutdown request immediately transitions stimulator to SAFE_OFF state and halts pulse discharge.",
        "TC-404: Verify that therapy cancellation request transitions session status to CANCELLED and halts pending pulse sequence.",
        "TC-405: Verify that FFT spectral calculation algorithm accurately computes band powers across delta, theta, alpha, and beta neural frequencies.",
        "TC-406: Verify that resubmitting identical neural telemetry packet is rejected by telemetry deduplication trigger.",
        "TC-407: Verify that calibration audit query returns complete historical chronological trail of neurosurgeon voltage parameter adjustments.",
        "TC-408: Verify that biometric templates are stored as salted Argon2id hashes and no plaintext or reversible DES templates exist in database.",
        "TC-409: Verify that C++ closed-loop stimulation engine computes stimulation decisions within 20 milliseconds under 100 concurrent streams.",
        "TC-410: Verify that submitting clinical adverse event report with EDF recordings creates an open review ticket in safety committee queue.",
        "TC-411: Verify that neurology staff breakroom microwave oven heats lunches evenly without sparking.",
        "TC-412: Verify that surgical control room motorized ergonomic standing desks adjust height smoothly.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_NeuroFlow.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — NeuroFlow BCI Telemetry", [
        "CR-501: Enhance FR-102 to support high-frequency oscillation (HFO) biomarker detection in addition to spike-wave discharge synchrony.",
        "CR-502: Modify FR-109 to prohibit reversible DES encryption and mandate salted Argon2id hashing across all biometric clinical databases.",
        "CR-503: Extend FR-105 to calculate cross-frequency phase-amplitude coupling (PAC) between theta and gamma oscillations in hippocampal implants.",
        "CR-504: Enhance FR-104 to support automated SMS notifications to attending clinical fellows when a patient cancels neuromodulation therapy.",
        "CR-505: Update FS-203 to support gradual voltage ramping during emergency shutoff to prevent sudden neural rebound seizures.",
        "CR-506: Procure and install 25 motorized ergonomic standing desks for clinical EEG telemetry observation center.",
        "CR-507: Add cafeteria hot lunch weekly ordering module for surgical staff.",
        "CR-508: Add implant battery solar inductive charging carbon footprint dashboard.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_NeuroFlow.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — NeuroFlow BCI Telemetry", [
        "DEC-601: Medical board agreed that FR-106 telemetry deduplication must check sequence timestamps in RingBuffer memory cache.",
        "DEC-602: Cybersecurity director confirmed that FR-109 reversible DES encryption violates FDA premarket cybersecurity guidance and must be eliminated.",
        "DEC-603: Neurosurgery team discussed optogenetic neural stimulation probes but clinical human trial safety approval was not determined.",
        "DEC-604: Compliance director confirmed that FR-107 calibration audit logs must be retained for 25 years per FDA medical device regulations.",
        "DEC-605: Ethics committee confirmed that FR-108 dual-authorization rules must be enforced before adjusting stimulation voltage above 3.5 volts.",
        "DEC-606: QA lead agreed that TC-409 latency benchmarks must evaluate 100 concurrent patient streams on real-time Linux kernel staging node.",
        "DEC-607: Operations director decided to procure 25 motorized standing desks for neuro-telemetry monitoring staff.",
        "DEC-608: Science committee discussed quantum neural magnetometer arrays for non-contact magnetoencephalography but feasibility study is pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_NeuroFlow.docx"))

    print(f"[NeuroFlow] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
