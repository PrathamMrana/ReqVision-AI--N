"""
Generate 7 EduPay University Fee & Payment Platform documents.

Domain: Student tuition billing, course fees, payment gateways, receipts,
refund processing, bank reconciliation, bursar approvals, ledger audits.

Intentionally unmapped artifacts:
  US-310: "Export graduation gown rental receipts to microfiche for university archives"
  TC-410: "Verify cafeteria vending machine accepts coin currency"
  CR-506: "Install classroom smart projector scheduling system"
  DEC-607: "Procure 50 ergonomic chairs for the student union lounge"

Run: python generate_edupay_docs.py
Output: tests/edupay_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "edupay_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — EduPay University Billing Platform", [
        "BR-001: The EduPay platform shall provide a student fee portal where students can view their semester tuition invoices and outstanding course fee balances.",
        "BR-002: The system shall enable students and sponsors to complete fee payments using credit cards, debit cards, and net banking.",
        "BR-003: Finance staff shall be able to issue full or partial tuition refunds for approved course drops and semester withdrawals.",
        "BR-004: The bursar module shall generate official timestamped payment receipts and downloadable PDF tax certificates for all completed transactions.",
        "BR-005: The system shall dispatch payment confirmation receipts and overdue installment notifications to students via email and SMS.",
        "BR-006: Role-based authorization shall enforce separation of duties between fee assessment staff, bursar cashiers, and finance audit managers.",
        "BR-007: An immutable financial audit trail shall record every fee adjustment, refund authorization, and payment reversal for statutory compliance.",
        "BR-008: The platform shall reconcile daily bank settlement files against internal payment records and highlight discrepancy exceptions.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_EduPay.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — EduPay University Billing Platform", [
        "FR-101: The student dashboard shall retrieve tuition fee schedules, lab fees, and housing charges from the student ledger and display itemized fee breakdowns.",
        "FR-102: The payment processing engine shall integrate with external payment gateways (Stripe and Razorpay) to process credit card, debit card, and net banking transactions.",
        "FR-103: The refund management service shall allow authorized finance officers to calculate and disburse refunds for dropped courses according to the university refund schedule.",
        "FR-104: The receipt generation service shall create verifiable PDF payment receipts with unique transaction reference numbers and downloadable tax statements.",
        "FR-105: The notification service shall dispatch email and SMS alerts for successful payment confirmations, upcoming installment due dates, and overdue fee warnings.",
        "FR-106: The access control subsystem shall enforce RBAC permissions preventing cashiers from approving refunds and restricting students to their own financial records.",
        "FR-107: The financial audit service shall write an immutable append-only record for all fee adjustments, refund approvals, and transaction state changes to the audit ledger.",
        "FR-108: The automated reconciliation engine shall ingest daily bank settlement CSV files and match transaction identifiers against internal ledger entries.",
        "FR-109: Student password credentials shall be stored using reversible DES encryption so that the student helpdesk can read original passwords to students over the phone.",
        "FR-110: The platform shall support a minimum throughput of 300 concurrent payment checkouts per minute during peak semester registration rush hours.",
        "NFR-101: The payment subsystem shall comply with PCI-DSS Level 1 security standards for credit card handling and tokenization.",
        "NFR-102: The fee payment portal shall maintain 99.95% availability during the scheduled two-week semester fee payment window.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_EduPay.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — EduPay University Billing Platform", [
        "FS-201: Implement student fee schedule API endpoint GET /api/students/{id}/fees returning itemized tuition, lab, and housing line items from fee_schedules table.",
        "FS-202: Implement payment gateway adapter integrating Stripe Charges API and Razorpay Webhooks; process card/UPI payments and update payment_transactions table.",
        "FS-203: Implement refund processing module with POST /api/refunds/disburse; validate approval token from finance manager; reverse payment via gateway API.",
        "FS-204: Implement PDF receipt generator using reportlab; embed QR code with cryptographic HMAC signature for receipt authenticity verification.",
        "FS-205: Implement notification worker listening on RabbitMQ queue; dispatch templated SMTP emails and Twilio SMS for transaction receipts and overdue reminders.",
        "FS-206: Implement RBAC middleware with JWT token claims validation; role hierarchy: BursarAdmin > FinanceOfficer > Cashier > Student; return 403 Forbidden for unauthorized endpoints.",
        "FS-207: Implement audit_ledger PostgreSQL table with append-only database triggers; record actor_id, action_type, entity_id, prev_balance, new_balance, timestamp.",
        "FS-208: Implement daily bank reconciliation batch job; parse MT940 / CSV bank settlement feeds; flag unmatched transactions in reconciliation_discrepancies table.",
        "FS-209: Student passwords shall be secured using argon2id one-way cryptographic hashing with per-user salts. Plaintext or reversible storage of passwords is strictly prohibited.",
        "FS-210: Implement redis-backed distributed rate limiter and horizontal pod autoscaler targeting 300 checkouts/minute at p99 latency under 1.5 seconds.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_EduPay.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — EduPay University Billing Platform", [
        "US-301: As a student, I want to view my current semester fee breakdown online so that I know the exact amount owed before the payment deadline.",
        "US-302: As a student, I want to pay my tuition fees using credit card or net banking so that I can settle my university dues securely from home.",
        "US-303: As a finance officer, I want to process approved refund requests for students who withdrew from courses so that money is returned to their original payment method.",
        "US-304: As a student, I want to download an official PDF receipt immediately after paying so that I have proof of tuition payment for scholarship and tax purposes.",
        "US-305: As a student, I want to receive an SMS and email notification when my fee payment is processed so that I have immediate confirmation.",
        "US-306: As a bursar auditor, I want an immutable audit trail of all manual fee waivers and refunds so that financial compliance can be audited.",
        "US-307: As a finance reconciliation clerk, I want an automated bank settlement comparison report so that discrepancies between bank deposits and student accounts are resolved.",
        "US-308: As a student, I want to log into the EduPay portal using my university email and password to access my financial account.",
        "US-309: As a systems administrator, I want the fee platform to handle peak fee registration rush hours without server timeouts.",
        "US-310: As a university archivist, I want to export all graduation gown rental receipts to microfiche format for physical storage in the campus basement vault.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_EduPay.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — EduPay University Billing Platform", [
        "TC-401: Verify that a student can retrieve their itemized semester tuition and lab fee balance, and amounts match the approved university fee schedule.",
        "TC-402: Verify that a student can complete a credit card tuition payment via Stripe, the transaction record is created, and the balance updates to zero.",
        "TC-403: Verify that a finance officer can disburse a tuition refund for an eligible dropped course, and the refund amount is credited to the student's card.",
        "TC-404: Verify that an official PDF receipt with verifiable QR code is generated upon payment completion and can be downloaded by the student.",
        "TC-405: Verify that payment confirmation SMS and email are delivered to the registered student contact within 10 seconds of payment completion.",
        "TC-406: Verify that attempting to UPDATE or DELETE an entry in the audit_ledger table is rejected by PostgreSQL database triggers with a permission error.",
        "TC-407: Verify that the automated reconciliation batch job correctly identifies and flags an unrecorded bank credit in the reconciliation_discrepancies table.",
        "TC-408: Verify that user passwords are stored as argon2id salted hashes and that no plaintext or reversible passwords exist in the user credentials table.",
        "TC-409: Verify that the checkout service sustains 300 concurrent payment transactions per minute with response time under 1.5 seconds during load testing.",
        "TC-410: Verify that the campus cafeteria vending machine dispenses chocolate bars when students insert coins — this test is for the canteen snack vendor.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_EduPay.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — EduPay University Billing Platform", [
        "CR-501: Enhance FR-102 to support Apple Pay and Google Pay digital wallet checkouts in addition to standard credit card and net banking.",
        "CR-502: Modify FR-109 to prohibit reversible credential encryption and mandate salted argon2id password hashing in compliance with ISO 27001.",
        "CR-503: Update FR-105 to support WhatsApp payment receipts and installment alert notifications in addition to SMS and email.",
        "CR-504: Extend FR-110 to support 600 concurrent payment checkouts per minute for the university wide enrollment day expansion.",
        "CR-505: Update FS-204 to include university official tax identification number and digital signature on all downloadable PDF payment receipts.",
        "CR-506: Install a classroom smart projector scheduling system in the engineering lecture hall to manage audio-visual equipment reservations.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_EduPay.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — EduPay University Billing Platform", [
        "DEC-601: Committee agreed that FR-101 fee schedule must support installment payment plans (3 equal split payments per semester) for undergraduate students.",
        "DEC-602: Security officer decided that FR-109 password recovery practice must be revoked immediately and replaced with secure self-service email token reset.",
        "DEC-603: Finance director confirmed that FR-108 automated bank reconciliation must execute daily at 02:00 AM UTC against the central bank SFTP server.",
        "DEC-604: Bursar team agreed that FR-104 receipts must retain a minimum 7-year retention period in compliance with national financial audit regulations.",
        "DEC-605: Architecture team confirmed that FR-106 RBAC permissions must be audited every quarter by an independent cybersecurity reviewer.",
        "DEC-606: QA team agreed that TC-402 payment gateway tests must be run against Stripe test mode sandbox with mock card numbers.",
        "DEC-607: Campus facilities team decided to procure 50 ergonomic chairs and 10 coffee tables for the student union lounge renovation.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_EduPay.docx"))

    print(f"[EduPay] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
