"""
Generate 7 TravelOps Corporate Travel & Expense Management documents.

Domain: Corporate trip booking, travel cost estimation, manager approval workflow,
pending trip request cancellation/withdrawal, mobile receipt OCR, duplicate receipt fraud check,
RBAC access control, finance expense reconciliation, concurrency scaling.

Intentionally unmapped/adversarial artifacts:
  US-310: "Export paper travel receipts to microfiche format for the basement physical archive"
  TC-410: "Verify airport executive lounge espresso coffee machine dispenses hot beverages"
  CR-506: "Add biometric international passport storage vault in travel agency partner office"
  DEC-603: "Team discussed faster approvals but did not define the meaning or implementation"
  DEC-607: "Procure 40 ergonomic mesh chairs for corporate travel coordinators"

Run: python generate_travelops_docs.py
Output: tests/travelops_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "travelops_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — TravelOps Corporate Travel & Expense", [
        "BR-001: The TravelOps platform shall provide an online portal where corporate employees can create and submit business travel requests.",
        "BR-002: The system shall calculate automated travel cost estimates including airfare, lodging, and per diem allowances before booking.",
        "BR-003: Designated managers shall review and approve or reject travel requests submitted by direct reports within the corporate hierarchy.",
        "BR-004: Employees shall be able to withdraw or cancel a pending travel request before manager approval has occurred.",
        "BR-005: The mobile application shall capture and upload digital photos of expense receipts for itemized claim reimbursement.",
        "BR-006: The expense engine shall automatically detect duplicate receipts and prevent duplicate expense claims from being submitted.",
        "BR-007: Role-based authorization shall enforce strict separation between corporate travelers, line managers, and finance auditors.",
        "BR-008: Finance staff shall reconcile approved expense claims against corporate credit card feeds and bank settlement records.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_TravelOps.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — TravelOps Corporate Travel & Expense", [
        "FR-101: The trip booking service shall allow employees to enter destination, travel dates, project codes, and submit travel requests.",
        "FR-102: The estimation engine shall compute projected trip costs based on corporate travel policies and preferred vendor rate tables.",
        "FR-103: The manager workflow service shall route pending travel requests to the employee's designated manager for approval or rejection.",
        "FR-104: The cancellation module shall permit employees to withdraw pending unapproved travel requests and release hold reservations.",
        "FR-105: The mobile receipt service shall allow employees to photograph receipts, perform OCR text extraction, and attach receipts to claims.",
        "FR-106: The fraud prevention module shall compute perceptual image hashes to detect and reject duplicate receipt uploads across claims.",
        "FR-107: The access control engine shall enforce RBAC permissions preventing employees from accessing other travelers' expense records.",
        "FR-108: The financial reconciliation service shall match approved expense line items against corporate credit card statement transactions.",
        "FR-109: Employee password credentials shall be stored using reversible DES encryption so that corporate IT helpdesk can recover passwords.",
        "FR-110: The platform shall support a minimum of 8000 concurrent month-end sessions during peak global expense submission rush.",
        "NFR-101: The mobile and web applications shall comply with SOC-2 Type II financial data security standards.",
        "NFR-102: The travel approval workflow shall maintain 99.95% system uptime during standard business operating hours.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_TravelOps.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — TravelOps Corporate Travel & Expense", [
        "FS-201: Implement trip submission REST API POST /api/trips; validate project budget allocation and create record in travel_requests table.",
        "FS-202: Implement travel cost estimation worker querying Sabre / Amadeus GDS fare matrix and corporate negotiated hotel rate APIs.",
        "FS-203: Implement manager approval workflow with JWT one-click email approval links and manager dashboard status action buttons.",
        "FS-204: Implement trip cancellation endpoint POST /api/trips/{id}/withdraw; transition status from PENDING to WITHDRAWN and notify manager.",
        "FS-205: Implement mobile receipt capture API POST /api/receipts/upload; invoke Tesseract OCR worker for total amount and date parsing.",
        "FS-206: Implement duplicate receipt detection using dHash perceptual image hashing and exact merchant amount comparison index.",
        "FS-207: Implement RBAC JWT security interceptor validating role claims: CorporateTraveler < DepartmentManager < FinanceAuditor < SysAdmin.",
        "FS-208: Implement bank and corporate card reconciliation engine parsing OFX / CSV feeds and matching expense transaction reference IDs.",
        "FS-209: Employee passwords shall be protected using salted PBKDF2 cryptographic hashing. Reversible or plaintext password storage is strictly forbidden.",
        "FS-210: Implement Kubernetes horizontal pod autoscaling with Redis session replication to sustain 8000 concurrent month-end sessions.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_TravelOps.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — TravelOps Corporate Travel & Expense", [
        "US-301: As a corporate employee, I want to submit a business travel request with itinerary dates so that my trip can be approved.",
        "US-302: As an employee, I want to see an automated cost estimate for my trip before submission so that I stay within corporate budget.",
        "US-303: As a department manager, I want to review and approve travel requests from my team so that travel expenditure is authorized.",
        "US-304: As an employee, I want to withdraw a pending travel request that is no longer needed so that budget holds are released.",
        "US-305: As a business traveler on the go, I want to snap a photo of my meal receipt on my phone so that it is attached to my expense report.",
        "US-306: As an employee, I want to log into my TravelOps account using my corporate email and secure password to view my trip history.",
        "US-307: As a finance compliance officer, I want an immutable audit log of all expense approvals and reimbursement payouts for statutory audits.",
        "US-308: As a finance reconciliation clerk, I want to reconcile approved expense claims against corporate credit card feeds automatically.",
        "US-309: As a corporate traveler, I want the expense portal to remain responsive without lag during month-end expense filing deadlines.",
        "US-310: As a corporate archivist, I want to export paper travel receipts to microfiche format for physical storage in the basement archive vault.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_TravelOps.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — TravelOps Corporate Travel & Expense", [
        "TC-401: Verify that an employee can submit a business travel request, the record is stored in travel_requests table, and status is PENDING.",
        "TC-402: Verify that automated cost estimation calculates correct airfare and per diem totals according to corporate policy rate tables.",
        "TC-403: Verify that a manager can approve a pending travel request, status updates to APPROVED, and notification is sent to employee.",
        "TC-404: Verify that an employee can withdraw a pending travel request, status transitions to WITHDRAWN, and booking holds are released.",
        "TC-405: Verify that uploading a mobile receipt photo extracts merchant name and amount correctly and associates receipt with expense claim.",
        "TC-406: Verify that uploading a duplicate receipt image flags the claim with a duplicate warning and blocks duplicate reimbursement.",
        "TC-407: Verify that attempting to access another employee's expense claim without manager or auditor permissions returns 403 Forbidden.",
        "TC-408: Verify that user passwords are stored as salted PBKDF2 hashes and no reversible or plaintext passwords exist in credentials table.",
        "TC-409: Verify that the platform sustains 8000 concurrent user sessions with p95 response time under 1.5s during month-end load testing.",
        "TC-410: Verify that the airport executive lounge espresso coffee machine dispenses hot beverages when pressing the cappuccino button.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_TravelOps.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — TravelOps Corporate Travel & Expense", [
        "CR-501: Enhance FR-103 to send manager travel approval notifications and action buttons via Slack in addition to email.",
        "CR-502: Modify FR-109 to prohibit reversible password storage and mandate salted PBKDF2 password hashing in compliance with SOC-2.",
        "CR-503: Extend FR-110 to support 15000 concurrent month-end sessions for the planned international subsidiary expansion.",
        "CR-504: Enhance FR-108 to support automated corporate American Express and Visa commercial card direct bank API reconciliation feeds.",
        "CR-505: Update FS-205 to support multi-currency receipt amount conversion using live European Central Bank foreign exchange rates.",
        "CR-506: Add a biometric international passport storage vault in travel agency partner office for physical visa stamp handling.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_TravelOps.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — TravelOps Corporate Travel & Expense", [
        "DEC-601: Committee agreed that FR-103 manager approvals must support Slack interactive messages for faster travel authorizations.",
        "DEC-602: Security lead decided that FR-109 password recovery practice must be revoked and replaced with self-service password reset.",
        "DEC-603: The team discussed faster approvals but did not define the meaning or implementation specifications for the system.",
        "DEC-604: Finance director confirmed that FR-108 bank reconciliation must run daily at 01:00 AM UTC against corporate bank SFTP.",
        "DEC-605: Architecture team confirmed that FR-107 RBAC permission rules must undergo mandatory quarterly access reviews.",
        "DEC-606: QA team agreed that TC-409 concurrency load test must be executed on a dedicated staging environment replicating production.",
        "DEC-607: Corporate facilities manager decided to procure 40 ergonomic mesh chairs for corporate travel coordinators.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_TravelOps.docx"))

    print(f"[TravelOps] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
