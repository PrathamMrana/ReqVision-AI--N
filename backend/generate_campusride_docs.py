import os
from docx import Document

os.makedirs('tests/campusride_docs', exist_ok=True)

# 1. 01_BRD_CampusRide.docx
doc1 = Document()
doc1.add_heading('Business Requirements Document - CampusRide Platform', 0)
doc1.add_paragraph('Business Objective: Deliver a comprehensive enterprise campus shuttle transit platform.')
doc1.add_paragraph('Stakeholders: Campus Transportation Director, Student Affairs, Campus Police, Fleet Manager.')
doc1.add_paragraph('BR-001: Reduce average campus shuttle wait time during peak class changes.') # Intentionally unmapped in SRS
doc1.add_paragraph('BR-002: The platform shall provide live arrival board and real-time shuttle ETA tracking for students.')
doc1.add_paragraph('BR-003: The system shall support seat reservation and booking for high-demand campus shuttle routes.')
doc1.add_paragraph('BR-004: The system must prevent unauthorized riders from booking staff-only routes.')
doc1.add_paragraph('BR-005: The platform shall compile operations analytics and ridership reports for fleet management.')
doc1.add_paragraph('BR-006: The system shall broadcast service alert notifications for shuttle delays and route diversions.')
doc1.add_paragraph('BR-007: The platform shall maintain an immutable audit ledger for all ride bookings and safety transitions.')
doc1.add_paragraph('BR-008: The platform shall provide wheelchair-friendly vehicle information and accessibility booking.')
doc1.save('tests/campusride_docs/01_BRD_CampusRide.docx')

# 2. 02_SRS_CampusRide.docx
doc2 = Document()
doc2.add_heading('Software Requirements Specification - CampusRide Transit System', 0)
doc2.add_paragraph('Scope: Campus shuttle software functional and non-functional engineering specifications.')
doc2.add_paragraph('FR-101: The system shall provide real-time campus shuttle ETA tracking and live arrival board with sub-200ms latency.')
doc2.add_paragraph('FR-102: The system shall process seat reservation and booking service for high-demand campus shuttle routes.')
doc2.add_paragraph('FR-103: The system shall enforce role-based authorization matrix restricting staff-only shuttle routes to verified faculty credentials.')
doc2.add_paragraph('FR-104: The service alert notification engine shall dispatch push and SMS alerts for shuttle delays and route diversions.')
doc2.add_paragraph('FR-105: The fleet operations analytics module shall compile ride metrics, passenger loads, and route utilization into reports.')
doc2.add_paragraph('FR-106: The system shall record all ride bookings, cancellations, and driver assignments in an immutable audit ledger.')
doc2.add_paragraph('FR-107: The vehicle accessibility module shall flag wheelchair boarding capacity and ramp availability for incoming shuttles.')
doc2.add_paragraph('FR-108: The seat reservation cancellation endpoint shall allow riders to cancel bookings prior to cutoff window.')
doc2.add_paragraph('FR-109: The capacity administration service shall allow fleet managers to configure shuttle vehicle passenger limits.')
doc2.add_paragraph('FR-110: The platform shall support scalable high throughput architecture handling concurrent rider requests during rush hour peaks.')
doc2.add_paragraph('NFR-101: System response time shall be under 500ms for real-time GPS coordinate updates.')
doc2.add_paragraph('NFR-102: The platform shall maintain 99.9% uptime availability for fleet dispatch services.')
doc2.add_paragraph('NFR-103: The system shall enforce data retention compliance for rider logs.')
doc2.save('tests/campusride_docs/02_SRS_CampusRide.docx')

# 3. 03_FRD_CampusRide.docx
doc3 = Document()
doc3.add_heading('Functional Requirements Document - CampusRide Subsystems', 0)
doc3.add_paragraph('Functional Specification: Component-level process flows and interface behaviors.')
doc3.add_paragraph('FS-201: Arrival board service indexes shuttle GPS feeds and computes real-time arrival countdowns.')
doc3.add_paragraph('FS-202: Seat reservation controller allocates seat inventory and generates digital boarding passes.')
doc3.add_paragraph('FS-203: Route authorization middleware validates university staff credentials on staff-designated shuttle routes.')
doc3.add_paragraph('FS-204: Push and SMS alert dispatch service transmits delay notifications and route diversions to subscribed riders.')
doc3.add_paragraph('FS-205: Analytics reporting engine aggregates shuttle ridership metrics into printable daily fleet dashboards.')
doc3.add_paragraph('FS-206: Ledger service persists immutable ride booking transitions to PostgreSQL audit table.')
doc3.add_paragraph('FS-207: Accessibility flag filter identifies wheelchair-accessible shuttles and allocates ramp priority spaces.')
doc3.add_paragraph('FS-208: Reservation cancellation handler releases seat quota when rider cancels before departure cutoff.')
doc3.add_paragraph('FS-209: Vehicle capacity configuration module allows administrators to adjust shuttle seat limits.')
doc3.add_paragraph('FS-210: Load balancing cluster distribution handling high throughput during campus rush hours.')
doc3.save('tests/campusride_docs/03_FRD_CampusRide.docx')

# 4. 04_User_Stories_CampusRide.docx
doc4 = Document()
doc4.add_heading('User Stories - CampusRide Sprint Backlog', 0)
doc4.add_paragraph('US-301: As a Student Rider, I want to view live shuttle arrival times so that I know when my shuttle reaches the stop.')
doc4.add_paragraph('US-302: As a Student Rider, I want to reserve a seat on the express shuttle so that I am guaranteed boarding.')
doc4.add_paragraph('US-303: As a University Staff Member, I want to book staff-only shuttle routes using my employee ID.')
doc4.add_paragraph('US-304: As a Rider, I want to receive instant alerts when my shuttle route is delayed or diverted.')
doc4.add_paragraph('US-305: As a Fleet Manager, I want to view daily ridership analytics so that I can optimize vehicle schedules.')
doc4.add_paragraph('US-306: As a Compliance Auditor, I want to inspect ride booking logs to ensure safety and regulatory adherence.')
doc4.add_paragraph('US-307: As a Rider with Mobility Needs, I want to check wheelchair ramp availability on incoming shuttles.')
doc4.add_paragraph('US-308: As a Rider, I want to cancel my active seat reservation before the departure cutoff time.')
doc4.add_paragraph('US-309: As a Transportation Administrator, I want to update vehicle capacity limits for special campus events.')
doc4.add_paragraph('US-310: As an Administrator, I want to export legacy campus vehicle registration records to magnetic tape.') # Intentionally unmapped
doc4.save('tests/campusride_docs/04_User_Stories_CampusRide.docx')

# 5. 05_Test_Cases_CampusRide.docx
doc5 = Document()
doc5.add_heading('Test Case Specification - CampusRide Verification Suite', 0)
doc5.add_paragraph('TC-401: Test Scenario: Live Arrival Board Updates. Verify shuttle ETA updates on arrival board with low latency.')
doc5.add_paragraph('TC-402: Test Scenario: Seat Reservation Flow. Verify reserving a seat decrements available capacity and issues boarding pass.')
doc5.add_paragraph('TC-403: Test Scenario: Staff-Route Authorization. Verify student credential cannot book staff-only express shuttle.')
doc5.add_paragraph('TC-404: Test Scenario: Service Alert Dispatch. Verify push notification is sent to riders when route delay occurs.')
doc5.add_paragraph('TC-405: Test Scenario: Fleet Analytics Dashboard. Verify daily ridership report generates accurate boarding totals.')
doc5.add_paragraph('TC-406: Test Scenario: Audit Ledger Persistence. Verify ride booking and cancellation events are logged to immutable table.')
doc5.add_paragraph('TC-407: Test Scenario: Wheelchair Accessibility Verification. Verify vehicle accessibility flag correctly displays ramp availability.')
doc5.add_paragraph('TC-408: Test Scenario: Reservation Cancellation. Verify canceling reservation before cutoff restores seat capacity.')
doc5.add_paragraph('TC-409: Test Scenario: Capacity Administration. Verify administrator adjusts vehicle maximum passenger limit.')
doc5.add_paragraph('TC-410: Test Scenario: Cafeteria Meal Plan Verification. Verify campus dining card barcode scan.') # Intentionally unmapped
doc5.save('tests/campusride_docs/05_Test_Cases_CampusRide.docx')

# 6. 06_Change_Requests_CampusRide.docx
doc6 = Document()
doc6.add_heading('Change Requests - CampusRide Engineering Change Collection', 0)
doc6.add_paragraph('CR-501: Optimize GPS coordinate refresh interval to 250ms for arrival board accuracy.')
doc6.add_paragraph('CR-502: Add audio voice announcements for service delay alerts in addition to push notifications.')
doc6.add_paragraph('CR-503: Allow designated graduate teaching assistants to access staff-only shuttle routes.')
doc6.add_paragraph('CR-504: Enforce 99.95% high availability clustering for dispatch servers during finals week.')
doc6.add_paragraph('CR-505: Add automated vehicle capacity overflow alerts for shuttle fleet managers.')
doc6.add_paragraph('CR-506: Introduce a driver shift-planning and overtime payroll module.') # Intentionally unmapped
doc6.save('tests/campusride_docs/06_Change_Requests_CampusRide.docx')

# 7. 07_Meeting_Minutes_CampusRide.docx
doc7 = Document()
doc7.add_heading('Meeting Minutes - CampusRide Architecture Review Board', 0)
doc7.add_paragraph('DEC-601: Transportation board discussed reducing peak campus shuttle wait time BR-001.')
doc7.add_paragraph('DEC-602: Infrastructure committee approved peak-hour traffic scaling FR-110.')
doc7.add_paragraph('DEC-603: Accessibility advisory board reviewed wheelchair boarding standards BR-008.')
doc7.add_paragraph('DEC-604: Fleet operations confirmed daily analytics dashboard BR-005.')
doc7.add_paragraph('DEC-605: Compliance team reviewed audit ledger security BR-007.')
doc7.add_paragraph('DEC-606: Product team approved audio announcements CR-502.')
doc7.add_paragraph('DEC-607: Team discussed replacing physical campus parking-gate barrier hardware.') # Intentionally unmapped
doc7.save('tests/campusride_docs/07_Meeting_Minutes_CampusRide.docx')

print('All 7 CampusRide test documents successfully generated in tests/campusride_docs/!')
