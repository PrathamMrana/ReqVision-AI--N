"""
Generate 7 FinVault Cloud Core Banking & Ledger Audit documents.

Domain: Double-entry core banking ledger, real-time ACH/wire payment settlement,
PCI-DSS cryptographic transaction tokenization, account balance holds, interest accruals,
regulatory liquidity reporting, customer dispute chargebacks.

Intentionally unmapped/adversarial artifacts:
  US-310: "As a retail bank customer, I want to view my carbon footprint score on monthly debit purchases."
  TC-410: "Verify that customer carbon emission score is calculated from debit transactions."
  US-311: "As a compliance archivist, I want to export SWIFT settlement certificates to 35mm microfiche film."
  CR-506: "Add instant biometric iris scanning for ATM cash withdrawals."
  CR-507: "Procure and install 40 ergonomic sit-stand desks for treasury operations floor."
  DEC-607: "Risk committee discussed automated crypto trading arbitrage but regulatory approval was not determined."

Run: python generate_finvault_docs.py
Output: tests/finvault_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "finvault_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — FinVault Core Banking", [
        "BR-001: The FinVault core banking platform shall maintain an immutable double-entry ledger for customer deposit accounts.",
        "BR-002: The system shall process real-time domestic interbank wire payments and automated clearing house transfers.",
        "BR-003: Bank customer service representatives shall be able to place temporary debit holds on suspicious accounts.",
        "BR-004: Customers shall be able to cancel or withdraw a scheduled future recurring bill payment before daily batch cutoff.",
        "BR-005: The banking platform shall calculate daily compound interest accruals on high-yield savings accounts.",
        "BR-006: The platform shall automatically detect and prevent duplicate payment transactions within a 60-second window.",
        "BR-007: Treasury officers shall maintain an immutable chronological audit trail of all manual ledger journal adjustments.",
        "BR-008: Role-based authorization shall enforce strict four-eyes compliance separation between loan underwriters and fund cashiers.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_FinVault.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — FinVault Core Banking", [
        "FR-101: The core ledger service shall record balanced debits and credits across customer deposit and bank liability accounts.",
        "FR-102: The payment gateway engine shall ingest ISO 20022 wire messages and dispatch settlement notifications to Federal Reserve.",
        "FR-103: The account hold module shall allow compliance staff to place administrative debit freezes on customer checking accounts.",
        "FR-104: The scheduled payment cancellation service shall allow customers to withdraw pending scheduled bill payments.",
        "FR-105: The interest computation engine shall calculate daily savings interest based on annualized percentage yield APY tiers.",
        "FR-106: The transaction validation engine shall detect and block duplicate payment requests sharing identical amount and reference.",
        "FR-107: The ledger audit service shall store an immutable queryable audit log of manual journal entries and supervisor approvals.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing branch tellers from approving credit limit increases.",
        "FR-109: Customer account passwords shall be stored using reversible DES encryption so that phone bankers can assist forgotten callers.",
        "FR-110: The dispute resolution module shall allow customers to submit transaction dispute claims and upload merchant receipts.",
        "NFR-101: The core ledger posting API shall achieve p99 transaction write latency under 50 milliseconds under 10000 concurrent TPS.",
        "NFR-102: Core banking settlement ledger services shall maintain 99.999% high availability with zero data loss RPO.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_FinVault.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — FinVault Core Banking", [
        "FS-201: Implement double-entry ledger posting REST API POST /api/ledger/entries recording balanced debit and credit legs in PostgreSQL.",
        "FS-202: Implement ISO 20022 wire payment processor worker ingesting PACS.008 wire messages and updating settlement status.",
        "FS-203: Implement account hold administration module transitioning account status to DEBIT_FROZEN and blocking card auths.",
        "FS-204: Implement recurring payment cancellation endpoint POST /api/payments/recurring/{id}/cancel; transition status to CANCELLED.",
        "FS-205: Implement savings interest accrual nightly batch job computing daily yield and crediting monthly interest payouts.",
        "FS-206: Implement duplicate transaction filter trigger rejecting identical account, amount, and counterparty requests with HTTP 409.",
        "FS-207: Implement historical ledger adjustment query API GET /api/ledger/adjustments returning immutable audit trails.",
        "FS-208: Implement RBAC JWT token validator enforcing Teller < BranchManager < ChiefRiskOfficer authorization hierarchies.",
        "FS-209: Customer account passwords shall be secured using salted Argon2id cryptographic hashing. Reversible DES is strictly prohibited.",
        "FS-210: Implement distributed Go ledger posting engine with RocksDB in-memory journal buffers to achieve p99 latency under 50ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_FinVault.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — FinVault Core Banking", [
        "US-301: As an account holder, I want to transfer funds between my checking and savings accounts so that money is reallocated.",
        "US-302: As a business customer, I want to initiate an outbound wire transfer so that vendors receive funds on the same day.",
        "US-303: As a compliance officer, I want to place a debit freeze on a compromised account so that unauthorized withdrawals are prevented.",
        "US-304: As a retail banking customer, I want to cancel an un-executed scheduled bill payment so that my account is not debited.",
        "US-305: As a savings customer, I want to view daily accumulated interest earnings on my high-yield savings account.",
        "US-306: As a customer, I want the banking app to prevent duplicate payments if I tap submit twice during slow connectivity.",
        "US-307: As an internal bank auditor, I want to inspect historical ledger journal adjustments and supervisor approvals.",
        "US-308: As a customer service agent, I want to log into the CRM portal with my multi-factor credentials to assist customers.",
        "US-309: As a mobile banking customer during payday rush, I want the banking app to load balances quickly without freezing.",
        "US-310: As a retail bank customer, I want to view my estimated carbon footprint score on monthly debit purchases.",
        "US-311: As a compliance archivist, I want to export SWIFT settlement certificates to 35mm microfiche film for national library archives.",
        "US-312: As an account holder, I want to submit a transaction dispute and attach digital photos of merchant receipts.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_FinVault.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — FinVault Core Banking", [
        "TC-401: Verify that posting a ledger entry creates debit and credit records that sum to zero in ledger_entries table.",
        "TC-402: Verify that incoming ISO 20022 wire transfer packet updates recipient account balance and dispatches confirmation.",
        "TC-403: Verify that placing account freeze blocks subsequent debit card transactions with account restricted response.",
        "TC-404: Verify that scheduled payment cancellation transitions payment status to CANCELLED and cancels queue job.",
        "TC-405: Verify that nightly interest calculation correctly applies APY rate to minimum daily ledger balance.",
        "TC-406: Verify that resubmitting identical payment request within 60 seconds is rejected by duplicate filter trigger.",
        "TC-407: Verify that ledger audit query returns complete historical chronological trail of supervisor journal overrides.",
        "TC-408: Verify that customer passwords are stored as salted Argon2id hashes and no reversible or plaintext passwords exist in DB.",
        "TC-409: Verify that distributed Go ledger posting engine achieves p99 latency under 50ms under 10000 concurrent TPS.",
        "TC-410: Verify that submitting transaction dispute with receipt photos creates a chargeback case in operations queue.",
        "TC-411: Verify that branch breakroom espresso coffee machine powers on and brews correctly.",
        "TC-412: Verify that treasury sit-stand motorized desks adjust height smoothly.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_FinVault.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — FinVault Core Banking", [
        "CR-501: Enhance FR-102 to support FedNow instant 24/7 payment settlement rail in addition to traditional wire transfers.",
        "CR-502: Modify FR-109 to prohibit reversible DES encryption and mandate salted Argon2id password hashing across all channels.",
        "CR-503: Extend FR-105 to support tiered interest rates for high-net-worth commercial treasury balances exceeding 5 million dollars.",
        "CR-504: Enhance FR-104 to support automated SMS and email notifications to account holders when scheduled payment is cancelled.",
        "CR-505: Update FS-203 to support partial debit holds allowing recurring utility debits while freezing external wire transfers.",
        "CR-506: Add instant biometric iris scanning authentication module for ATM cash withdrawal kiosks.",
        "CR-507: Procure and install 40 ergonomic motorized sit-stand desks for treasury operations trading floor.",
        "CR-508: Add retail customer carbon footprint scoring dashboard based on debit card merchant categories.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_FinVault.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — FinVault Core Banking", [
        "DEC-601: Architecture committee agreed that FR-106 duplicate payment prevention must check idempotency keys in Redis cache.",
        "DEC-602: Security officer confirmed that FR-109 reversible DES encryption is a severe compliance violation and must be replaced.",
        "DEC-603: Risk committee discussed automated crypto trading arbitrage but regulatory approval and scope were not determined.",
        "DEC-604: Compliance director confirmed that FR-107 historical ledger audit trails must be retained for 10 years for SEC compliance.",
        "DEC-605: Audit committee confirmed that FR-108 RBAC dual-control rules must be enforced for all wire transfers exceeding 50000 dollars.",
        "DEC-606: QA team agreed that TC-409 load tests must benchmark 10000 TPS on AWS staging cluster before production rollout.",
        "DEC-607: Operations director decided to procure 40 ergonomic sit-stand motorized desks for treasury settlement staff.",
        "DEC-608: Product team discussed quantum-resistant encryption algorithms but technical feasibility review is pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_FinVault.docx"))

    print(f"[FinVault] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
