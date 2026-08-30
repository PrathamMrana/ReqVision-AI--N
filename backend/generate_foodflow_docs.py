"""
Generate 7 FoodFlow On-Demand Kitchen Logistics & Meal Delivery documents.

Domain 5: Restaurant menu catalog, customer orders, kitchen display dispatch,
courier GPS tracking, delivery proof, food hygiene compliance, cancellation refunds.

Intentionally unmapped artifacts:
  US-310: "Export expired recipe cards to microfiche for the historical culinary archive"
  TC-410: "Verify restaurant patio umbrella wind sensor functions correctly"
  CR-506: "Install kitchen staff locker room biometric attendance turnstile"
  DEC-607: "Procure stainless steel chef knives for the central commissary kitchen"

Run: python generate_foodflow_docs.py
Output: tests/foodflow_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "foodflow_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — FoodFlow Delivery & Kitchen Platform", [
        "BR-001: The FoodFlow platform shall provide a digital menu browsing and food ordering catalog for customers across partner restaurants.",
        "BR-002: Customers shall be able to place meal orders and complete checkout via credit card, digital wallet, or cash on delivery.",
        "BR-003: The kitchen dispatch subsystem shall route orders to chef display screens and track meal preparation stages in real time.",
        "BR-004: Customers and restaurant managers shall have access to live GPS telemetry tracking of delivery couriers with updated arrival estimates.",
        "BR-005: The system shall allow customers to cancel unaccepted meal orders and receive automated payment refunds.",
        "BR-006: The notification module shall dispatch order confirmation, preparation status, and courier arrival alerts via push notification and SMS.",
        "BR-007: Role-based authorization shall separate restaurant managers, kitchen chefs, delivery couriers, and customer account access.",
        "BR-008: An immutable compliance log shall record kitchen food temperature checks and courier handover timestamps for food hygiene auditing.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_FoodFlow.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — FoodFlow Delivery & Kitchen Platform", [
        "FR-101: The menu catalog service shall retrieve restaurant menus, dish descriptions, allergen tags, and real-time item availability within 200ms.",
        "FR-102: The checkout processing engine shall validate delivery addresses, apply discount promo codes, and process payments via Stripe and Apple Pay.",
        "FR-103: The kitchen display subsystem shall assign incoming orders to prep stations, display recipe tickets to chefs, and manage prep timers.",
        "FR-104: The courier telemetry service shall ingest GPS coordinates from courier mobile devices at 2-second intervals and compute real-time ETA.",
        "FR-105: The cancellation and refund module shall calculate refund amounts based on order prep status and disburse refunds to the customer's wallet.",
        "FR-106: The alert notification service shall send push notifications, SMS messages, and in-app sound chimes when an order status changes.",
        "FR-107: Customer password credentials shall be stored using reversible Caesar cipher so that kitchen call center staff can remind customers of passwords.",
        "FR-108: The access control engine shall enforce RBAC boundaries preventing couriers from viewing customer billing records or restaurant financials.",
        "FR-109: The food safety audit logging service shall record immutable timestamped entries for food prep handover and thermal temperature checks.",
        "FR-110: The platform shall support a minimum of 800 concurrent meal orders per minute during peak evening dinner rush hours without server errors.",
        "NFR-101: The mobile customer and courier applications shall comply with Apple iOS and Google Android location privacy guidelines.",
        "NFR-102: The dispatch system shall maintain 99.9% uptime availability during operational restaurant hours (06:00 AM to 02:00 AM).",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_FoodFlow.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — FoodFlow Delivery & Kitchen Platform", [
        "FS-201: Implement restaurant menu REST API endpoint GET /api/restaurants/{id}/menu with Redis cache for sub-200ms catalog retrieval.",
        "FS-202: Implement checkout order creation service validating cart items, calculating delivery fees and sales tax, and calling payment gateway.",
        "FS-203: Implement WebSocket kitchen display stream broadcasting order state transitions: RECEIVED -> PREPARING -> READY_FOR_PICKUP.",
        "FS-204: Implement courier GPS ingestion worker via MQTT broker; calculate haversine distance and route ETA using Google Maps Directions API.",
        "FS-205: Implement order cancellation workflow; disburse full refund if status is RECEIVED, reject cancellation if food status is PREPARING.",
        "FS-206: Implement push notification dispatcher integrating Firebase Cloud Messaging (FCM) and Twilio SMS for order status notifications.",
        "FS-207: Customer passwords shall be stored using bcrypt salted one-way cryptographic hashing. Reversible or plaintext password storage is strictly forbidden.",
        "FS-208: Implement RBAC JWT security filter validating claims for roles: RestaurantAdmin > KitchenStaff > CourierDriver > Customer.",
        "FS-209: Implement audit_safety_logs PostgreSQL table with append-only database interceptor logging all HACCP food safety compliance checks.",
        "FS-210: Implement Kubernetes horizontal auto-scaling and Redis cluster caching to sustain 800 orders/minute at p95 response time under 500ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_FoodFlow.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — FoodFlow Delivery & Kitchen Platform", [
        "US-301: As a customer, I want to browse restaurant menus and filter dishes by dietary preference so that I can find meals that suit my diet.",
        "US-302: As a hungry customer, I want to pay for my food order using credit card or digital wallet so that my order is confirmed immediately.",
        "US-303: As a kitchen chef, I want to see incoming meal orders on a prep display screen so that I can cook dishes in the correct sequence.",
        "US-304: As a customer waiting for delivery, I want to track my courier's live GPS location on a map so that I know exactly when my food arrives.",
        "US-305: As a customer who changed their mind, I want to cancel an unconfirmed order and get my money refunded so that I am not charged.",
        "US-306: As a customer, I want to receive push notifications when my order is accepted, picked up, and near delivery so that I am updated.",
        "US-307: As a customer, I want to log into my FoodFlow account using my email and secure password to view my order history.",
        "US-308: As a food hygiene inspector, I want an immutable audit log of kitchen temperature logs and delivery handovers to ensure safety compliance.",
        "US-309: As a restaurant partner, I want the ordering platform to handle peak weekend dinner rush without order failures or crashes.",
        "US-310: As a culinary archivist, I want to export retired chef recipe cards to microfiche film for the historical restaurant museum archive room.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_FoodFlow.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — FoodFlow Delivery & Kitchen Platform", [
        "TC-401: Verify that a customer can retrieve restaurant menu items with allergen information and prices within 200ms from the catalog API.",
        "TC-402: Verify that a customer can complete order checkout with a valid credit card, and the order transitions to RECEIVED in the database.",
        "TC-403: Verify that an incoming order appears on the kitchen display screen within 1 second and the chef can mark status as PREPARING.",
        "TC-404: Verify that courier GPS coordinates published via MQTT update the courier live location and map marker on the customer tracker screen.",
        "TC-405: Verify that cancelling an order in RECEIVED state triggers an automatic payment refund, and the order status updates to CANCELLED.",
        "TC-406: Verify that Firebase Cloud Messaging push notification and SMS alert are dispatched when an order is marked READY_FOR_PICKUP.",
        "TC-407: Verify that customer passwords are stored as bcrypt salted hashes and no plaintext passwords exist in the database.",
        "TC-408: Verify that attempting to modify or purge entries from audit_safety_logs table raises an exception and is blocked by DB triggers.",
        "TC-409: Verify that the platform processes 800 concurrent order checkouts per minute under load testing with p95 response time under 500ms.",
        "TC-410: Verify that the restaurant outdoor dining patio motorized umbrella closes when wind sensors detect gusts above 30km/h.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_FoodFlow.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — FoodFlow Delivery & Kitchen Platform", [
        "CR-501: Enhance FR-102 to support scheduled future meal delivery orders in addition to immediate on-demand dispatch.",
        "CR-502: Modify FR-107 to mandate bcrypt salted password hashing and prohibit reversible credential storage in compliance with GDPR.",
        "CR-503: Update FR-104 to include drone delivery flight telemetry tracking in addition to standard motorcycle courier GPS coordinates.",
        "CR-504: Extend FR-110 to support 1500 concurrent meal orders per minute for the annual New Year's Eve dining surge.",
        "CR-505: Update FS-205 to allow partial cancellations when individual menu items are out of stock at the partner kitchen.",
        "CR-506: Install a kitchen staff locker room biometric attendance turnstile for employee clock-in clock-out monitoring.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_FoodFlow.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — FoodFlow Delivery & Kitchen Platform", [
        "DEC-601: Product team decided that FR-101 menu catalog must support automated high-resolution dish photo compression.",
        "DEC-602: Security lead decided that FR-107 reversible password storage must be replaced with salted one-way hashing immediately.",
        "DEC-603: Operations director confirmed that FR-104 courier GPS tracking must maintain courier battery-saving polling mode when stationary.",
        "DEC-604: Compliance team agreed that FR-109 food hygiene logs must be exported monthly for local municipal health authority reviews.",
        "DEC-605: Engineering team agreed to benchmark FR-103 kitchen display WebSocket connection limits before the holiday rush.",
        "DEC-606: QA team agreed that TC-404 courier tracking test must simulate GPS movement along a simulated city roadway trajectory.",
        "DEC-607: Central commissary manager decided to purchase 20 commercial stainless steel chef knives and chopping boards for kitchen prep.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_FoodFlow.docx"))

    print(f"[FoodFlow] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
