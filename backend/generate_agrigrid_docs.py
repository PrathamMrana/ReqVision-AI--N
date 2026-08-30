"""
Generate 7 AgriGrid Autonomous Agricultural IoT Irrigation documents.

Domain: Precision agriculture, autonomous valve actuators, soil moisture telemetry,
nitrogen-phosphorus-potassium (NPK) nutrient dosing, evapotranspiration weather forecasting,
drone multispectral crop stress imaging, pesticide drift prevention.

Run: python generate_agrigrid_docs.py
Output: tests/agrigrid_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "agrigrid_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — AgriGrid Smart Irrigation", [
        "BR-001: The AgriGrid system shall maintain an automated irrigation schedule based on real-time soil moisture sensor telemetry.",
        "BR-002: The platform shall track weather forecast data to suppress automated irrigation cycles when precipitation is imminent.",
        "BR-003: Farm agronomists shall be able to place an emergency shutoff hold on compromised irrigation pump sectors.",
        "BR-004: Farm managers shall be able to cancel or abort a scheduled fertilizer fertigation cycle before valve actuation begins.",
        "BR-005: The system shall calculate soil nitrogen-phosphorus-potassium nutrient depletion rates based on crop growth stages.",
        "BR-006: The platform shall automatically detect and prevent duplicate sensor telemetry packet transmissions from solar field nodes.",
        "BR-007: Agricultural compliance inspectors shall inspect an immutable chronological audit trail of all manual pesticide spray overrides.",
        "BR-008: Role-based authorization shall enforce strict four-eyes compliance separation between field tractor operators and chemical agronomists.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_AgriGrid.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — AgriGrid Smart Irrigation", [
        "FR-101: The irrigation controller service shall actuate solenoid water valves when soil volumetric water content drops below 18 percent.",
        "FR-102: The meteorological ingestion service shall query radar precipitation APIs and defer active irrigation schedules by 24 hours.",
        "FR-103: The pump emergency shutdown module shall allow agronomists to transition pump status to EMERGENCY_STOP and depressurize mainlines.",
        "FR-104: The fertigation cancellation service shall allow farm operators to withdraw pending chemical nutrient injection cycles.",
        "FR-105: The soil nutrient model shall compute projected NPK consumption curves using thermal degree day accumulations and crop canopy size.",
        "FR-106: The telemetry ingestion engine shall detect and block duplicate sensor packets sharing identical device ID and sequence counter.",
        "FR-107: The spray audit service shall store an immutable queryable history of chemical spray approvals with GPS field boundary polygons.",
        "FR-108: The security engine shall enforce RBAC boundaries preventing irrigation technicians from authorizing restricted pesticide applications.",
        "FR-109: Agronomist portal passwords shall be stored using reversible Caesar encryption so that field supervisors can recover forgotten credentials.",
        "FR-110: The soil anomaly module shall allow field technicians to submit soil contamination incident tickets and attach drone multispectral photo exports.",
        "NFR-101: The field valve telemetry API shall dispatch actuator trigger commands within 300 milliseconds under 2000 concurrent sector events.",
        "NFR-102: The central farm management platform shall maintain 99.95% uptime during peak seasonal planting and harvest quarters.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_AgriGrid.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — AgriGrid Smart Irrigation", [
        "FS-201: Implement valve actuation REST API POST /api/sectors/{id}/irrigate sending LoRaWAN commands to field valve solenoids.",
        "FS-202: Implement weather radar worker querying NOAA precipitation feeds and transitioning sector schedule status to RAIN_DELAY.",
        "FS-203: Implement pump emergency shutdown endpoint POST /api/pumps/{id}/emergency-stop depressurizing mainline pressure relief valves.",
        "FS-204: Implement fertigation cancellation endpoint POST /api/fertigation/jobs/{id}/cancel; transition status to CANCELLED.",
        "FS-205: Implement nutrient computation batch job calculating NPK replenishment dosages based on sensor telemetry and soil lab test results.",
        "FS-206: Implement telemetry deduplication trigger rejecting redundant sensor packets with HTTP 409 Conflict status.",
        "FS-207: Implement spray audit query API GET /api/spray/logs returning immutable chronological audit records and agronomist digital signatures.",
        "FS-208: Implement RBAC JWT permission validator enforcing FieldWorker < IrrigationTech < LeadAgronomist authorization tiers.",
        "FS-209: Agronomist passwords shall be secured using salted Argon2id cryptographic hashing. Reversible Caesar encryption is strictly prohibited.",
        "FS-210: Implement distributed Go valve commanding service with Redis in-memory actuator buffers achieving command latency under 300ms.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_AgriGrid.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — AgriGrid Smart Irrigation", [
        "US-301: As a field technician, I want to trigger sector valve opening from my mobile tablet so that newly planted seedlings receive water.",
        "US-302: As a farm manager, I want the system to automatically pause irrigation during rainstorms so that water pumping costs are minimized.",
        "US-303: As an agronomist, I want to execute an emergency pump stop when a mainline pipe bursts so that soil erosion is prevented.",
        "US-304: As an irrigation operator, I want to cancel an unneeded fertilizer injection job so that chemical fertilizer is not wasted.",
        "US-305: As a crop consultant, I want to view calculated soil NPK depletion charts so that customized fertilizer recipes can be prepared.",
        "US-306: As a telemetry engineer, I want the sensor gateway to discard duplicate radio packet transmissions caused by LoRaWAN retries.",
        "US-307: As an EPA compliance auditor, I want to inspect chronological logs of chemical spray applications and agronomist supervisor sign-offs.",
        "US-308: As a farm security officer, I want to log into the management console using hardware security keys to protect irrigation controls.",
        "US-309: As an irrigation dispatcher during dry summer heatwave, I want the sector dashboard to refresh moisture levels quickly without lag.",
        "US-310: As a greenhouse worker, I want to view weekly cafeteria hot lunch menus in the farm staff portal.",
        "US-311: As an agricultural historian, I want to export seasonal crop harvest records to 35mm microfiche film for national library archives.",
        "US-312: As a soil scientist, I want to file a soil contamination ticket and attach drone multispectral photo exports.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_AgriGrid.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — AgriGrid Smart Irrigation", [
        "TC-401: Verify that POST /api/sectors/{id}/irrigate dispatches LoRaWAN open command to sector solenoid valve.",
        "TC-402: Verify that incoming NOAA precipitation forecast transitions scheduled sector irrigation to RAIN_DELAY status.",
        "TC-403: Verify that emergency pump shutdown depressurizes mainline pump and records emergency stop event in system log.",
        "TC-404: Verify that fertigation cancellation request transitions job status to CANCELLED and halts chemical metering pump.",
        "TC-405: Verify that NPK computation algorithm accurately calculates fertilizer replenishment requirements based on soil sensor readings.",
        "TC-406: Verify that resubmitting identical sensor packet within 30 seconds is rejected by telemetry deduplication trigger.",
        "TC-407: Verify that spray audit query returns complete historical chronological trail of agronomist chemical spray approvals.",
        "TC-408: Verify that agronomist passwords are stored as salted Argon2id hashes and no reversible Caesar passwords exist in database.",
        "TC-409: Verify that distributed Go valve commanding service executes trigger commands within 300 milliseconds under 2000 concurrent events.",
        "TC-410: Verify that submitting soil contamination ticket with drone photos creates an open investigation record in QA queue.",
        "TC-411: Verify that farm breakroom microwave oven heats lunches evenly without sparking.",
        "TC-412: Verify that agricultural greenhouse motorized standing desks adjust height smoothly.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_AgriGrid.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — AgriGrid Smart Irrigation", [
        "CR-501: Enhance FR-102 to ingest hyper-local on-farm acoustic disdrometer rain sensors in addition to regional radar APIs.",
        "CR-502: Modify FR-109 to prohibit reversible Caesar encryption and mandate salted Argon2id password hashing across all farm tablets.",
        "CR-503: Extend FR-105 to calculate soil micronutrient depletion for zinc, iron, and boron in commercial apple orchards.",
        "CR-504: Enhance FR-104 to support automated SMS notifications to field tractor operators when scheduled fertigation is aborted.",
        "CR-505: Update FS-203 to support partial sector isolation allowing drip zones to operate while emergency stopping pivot sprinklers.",
        "CR-506: Procure and install 30 motorized ergonomic sit-stand desks for farm dispatch operations center.",
        "CR-507: Add cafeteria hot lunch weekly ordering module for greenhouse harvest staff.",
        "CR-508: Add farm tractor carbon emission calculation dashboard based on diesel fuel consumption.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_AgriGrid.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — AgriGrid Smart Irrigation", [
        "DEC-601: Engineering committee agreed that FR-106 sensor deduplication must check telemetry sequence counters in Redis cache.",
        "DEC-602: Security officer confirmed that FR-109 reversible Caesar encryption violates farm cybersecurity standards and must be replaced.",
        "DEC-603: Agronomy team discussed autonomous crop-dusting drone swarm flight paths but regulatory airspace approval was not determined.",
        "DEC-604: Compliance director confirmed that FR-107 chemical spray audit logs must be retained for 15 years per Department of Agriculture rules.",
        "DEC-605: Safety committee confirmed that FR-108 dual-authorization rules must be enforced before applying restricted organophosphate pesticides.",
        "DEC-606: QA lead agreed that TC-409 load tests must benchmark 2000 concurrent valve events on staging environment before deployment.",
        "DEC-607: Operations director decided to procure 30 motorized standing desks for irrigation control room staff.",
        "DEC-608: Tech committee discussed quantum sensor magnetometers for deep root water sensing but feasibility research is pending.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_AgriGrid.docx"))

    print(f"[AgriGrid] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
