"""
Generate 7 Hospital Patient Management System documents for Phase 3 generalization testing.

Third domain — entirely different vocabulary from:
  - Online Library (books, borrowing, catalogue)
  - CampusRide (shuttle, transit, GPS, routes)

Hospital domain: patient, physician, appointment, prescription, medication,
lab result, billing, HIPAA, ward, diagnosis, admission, discharge, nurse, clinic.

Intentionally unmapped artifacts:
  US-310: "Export patient records to microfiche for legacy hospital archive"
  TC-410: No corresponding user story (test for system with no requirements)
  CR-506: "Install new cafeteria scheduling system in hospital canteen"
  DEC-607: "Procure new hospital furniture for ward renovation"

Run: python generate_health_docs.py
Output: tests/health_docs/
"""

import os
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "tests", "health_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_doc(title, content_lines):
    doc = Document()
    doc.add_heading(title, level=0)
    for line in content_lines:
        doc.add_paragraph(line)
    return doc


def main():
    # ── 01 BRD ───────────────────────────────────────────────────────────────
    brd = create_doc("Business Requirements Document — Hospital Patient Management System", [
        "BR-001: The hospital system shall provide a patient registration portal allowing new and returning patients to register their demographics, insurance, and medical history online.",
        "BR-002: The system shall allow patients to schedule, reschedule, and cancel appointments with physicians and clinic departments.",
        "BR-003: Physicians shall be able to create, update, and discontinue patient prescriptions with dosage and medication details recorded.",
        "BR-004: Patients and authorized physicians shall have secure access to lab results, diagnostic reports, and imaging records.",
        "BR-005: The billing subsystem shall generate itemized invoices for treatments, medications, and consultations, integrating with insurance claim processing.",
        "BR-006: The system shall send medication reminders, appointment alerts, and discharge notifications to patients via SMS and email.",
        "BR-007: Role-based access control shall restrict physician, nurse, administrator, and patient data access according to their assigned roles and ward assignments.",
        "BR-008: An immutable audit trail shall record all access and modifications to patient health records in compliance with HIPAA regulatory requirements.",
    ])
    brd.save(os.path.join(OUTPUT_DIR, "01_BRD_HospitalPMS.docx"))

    # ── 02 SRS ───────────────────────────────────────────────────────────────
    srs = create_doc("Software Requirements Specification — Hospital Patient Management System", [
        "FR-101: The system shall provide a patient registration interface supporting demographic data entry, insurance provider lookup, and medical history upload for new and existing patients.",
        "FR-102: The scheduling module shall allow patients to book appointments with physicians, select clinic departments, view available time slots, and receive confirmation within 30 seconds.",
        "FR-103: Physicians shall create and manage digital prescriptions including medication name, dosage, frequency, and duration; discontinued prescriptions shall be archived with reason.",
        "FR-104: Authorized users shall retrieve lab results, radiology reports, and pathology findings via a secure HTTPS portal with PDF download capability.",
        "FR-105: Patient credential storage shall use reversible encryption so that hospital administrators can recover original passwords when patients are locked out of the portal.",
        "FR-106: The notification service shall dispatch appointment reminders, medication alerts, and post-discharge care instructions via email and SMS to registered patient contacts.",
        "FR-107: The billing engine shall generate itemized patient invoices listing each procedure, medication, and consultation fee, and shall submit electronic claims to insurance providers.",
        "FR-108: The access control module shall enforce role-based permissions preventing nurses from accessing physician prescription records and restricting patient access to their own records only.",
        "FR-109: The audit logging service shall write an immutable timestamped record for every read, create, update, and delete operation on patient health records, stored in a separate audit table.",
        "FR-110: The system shall support a minimum of 500 concurrent authenticated users during peak morning admission hours without degradation of response time beyond 2 seconds.",
        "NFR-101: The system shall comply with HIPAA regulations for protected health information (PHI) storage, transmission, and access control.",
        "NFR-102: The system shall maintain 99.9% uptime availability, excluding scheduled maintenance windows of maximum 4 hours per month.",
    ])
    srs.save(os.path.join(OUTPUT_DIR, "02_SRS_HospitalPMS.docx"))

    # ── 03 FRD ───────────────────────────────────────────────────────────────
    frd = create_doc("Functional Requirements Document — Hospital Patient Management System", [
        "FS-201: Implement a patient registration REST API endpoint accepting demographic JSON payload; validate insurance provider against external claims API; store result in patient_registry table with encrypted PII fields.",
        "FS-202: Implement appointment scheduling service with PostgreSQL-backed availability calendar; enforce physician daily appointment limits; return confirmation token and notify patient via email within 30 seconds of booking.",
        "FS-203: Implement digital prescription management module with RESTful CRUD endpoints; prescriptions table with medication_name, dosage, frequency, duration, status, prescribing_physician_id; archive discontinued prescriptions with discontinuation_reason.",
        "FS-204: Implement lab results portal with secure JWT-authenticated GET endpoints; results fetched from lab_results table; PDF generation via reportlab; HTTPS-only transmission.",
        "FS-205: Patient passwords shall be stored using bcrypt salted one-way hashing. No plaintext or reversible credential storage is permitted. Password recovery shall use time-limited reset tokens sent to verified email addresses.",
        "FS-206: Implement notification dispatcher service supporting SMTP email and Twilio SMS channels; queue-based delivery with retry for appointment reminders, medication alerts, and discharge instructions.",
        "FS-207: Implement billing engine generating PDF itemized invoices; integrate with CMS-1500 electronic claim format for insurance submission; store invoice records in billing_records table.",
        "FS-208: Implement RBAC middleware validating JWT claims against role_permissions matrix; role hierarchy: Administrator > Physician > Nurse > Patient; route-level permission guards with 403 response for unauthorized access.",
        "FS-209: Implement audit_log PostgreSQL table with immutable append-only interceptor; log fields: user_id, role, action_type, entity_id, entity_type, timestamp, ip_address; no UPDATE or DELETE permitted on audit_log.",
        "FS-210: Implement horizontal auto-scaling via Kubernetes HPA; load balancer distributes across minimum 3 pod replicas; performance test at 500 concurrent users with p95 response time target of 2 seconds.",
    ])
    frd.save(os.path.join(OUTPUT_DIR, "03_FRD_HospitalPMS.docx"))

    # ── 04 User Stories ──────────────────────────────────────────────────────
    us = create_doc("User Stories — Hospital Patient Management System", [
        "US-301: As a new patient, I want to register my personal and insurance details online so that I can access hospital services without visiting the front desk.",
        "US-302: As a patient, I want to book an appointment with my physician and select a convenient time slot so that I can plan my hospital visit efficiently.",
        "US-303: As a physician, I want to prescribe medications digitally with dosage and frequency so that the pharmacy and patient have accurate treatment instructions.",
        "US-304: As a patient, I want to view my lab results and diagnostic reports securely online so that I can understand my health status without waiting for a physical copy.",
        "US-305: As a patient, I want to log into the hospital portal using my registered email and password so that I can securely access my health records.",
        "US-306: As a patient, I want to receive SMS and email reminders for upcoming appointments and medication schedules so that I do not miss my treatments.",
        "US-307: As an administrator, I want to generate itemized billing statements for patient treatments so that insurance claims can be submitted accurately.",
        "US-308: As a nurse, I want to access patient ward records and nursing notes while being prevented from modifying physician prescriptions so that clinical boundaries are maintained.",
        "US-309: As a compliance officer, I want to review an immutable audit log of all patient record accesses so that HIPAA compliance can be demonstrated during audits.",
        "US-310: As an IT archivist, I want to export all patient records to microfiche format for the legacy hospital archive room so that physical records are preserved in non-digital format.",
    ])
    us.save(os.path.join(OUTPUT_DIR, "04_User_Stories_HospitalPMS.docx"))

    # ── 05 Test Cases ────────────────────────────────────────────────────────
    tc = create_doc("Test Cases — Hospital Patient Management System", [
        "TC-401: Verify that a new patient can register with valid demographic and insurance data, and the record appears in the patient registry with all PII fields encrypted.",
        "TC-402: Verify that a patient can book an appointment with an available physician, receive a confirmation email within 30 seconds, and the appointment appears in the scheduling calendar.",
        "TC-403: Verify that a physician can create a prescription with medication name, dosage, frequency, and duration, and the prescription is visible to the patient and pharmacy.",
        "TC-404: Verify that a patient can retrieve their lab results via the HTTPS portal and download a PDF; verify that unauthorized users receive a 403 response.",
        "TC-405: Verify that patient passwords are stored as bcrypt hashes and that no plaintext passwords exist in the database; verify password reset uses email token flow.",
        "TC-406: Verify that appointment reminder SMS and email are dispatched 24 hours before appointment time and medication alerts are sent at scheduled intervals.",
        "TC-407: Verify that the billing engine generates a correctly itemized PDF invoice for a multi-procedure patient visit and submits an electronic claim to the insurance endpoint.",
        "TC-408: Verify that a nurse login cannot access physician prescription endpoints and receives a 403 Forbidden response; verify patient login cannot access other patients records.",
        "TC-409: Verify that every patient record access is written to the audit_log table; verify that attempting to UPDATE or DELETE an audit_log entry is rejected with an error.",
        "TC-410: Verify that the hospital tuck shop vending machine dispenses snacks correctly when staff insert coins — this test is for the canteen hardware system.",
    ])
    tc.save(os.path.join(OUTPUT_DIR, "05_Test_Cases_HospitalPMS.docx"))

    # ── 06 Change Requests ──────────────────────────────────────────────────
    cr = create_doc("Change Requests — Hospital Patient Management System", [
        "CR-501: Enhance FR-101 to support biometric patient identification (fingerprint scan) at registration kiosks in addition to online demographic form submission.",
        "CR-502: Modify FR-105 to align with HIPAA-compliant credential storage: require salted bcrypt hashing and prohibit reversible encryption for patient passwords.",
        "CR-503: Update FR-108 to add a new 'Radiologist' role with read-only access to imaging records and lab results, without access to billing or prescription data.",
        "CR-504: Extend FR-110 to support 2000 concurrent authenticated users during annual health camp peak periods by upgrading Kubernetes cluster scaling policy.",
        "CR-505: Update FS-204 to add real-time lab result push notifications when results are published, in addition to the existing portal retrieval mechanism.",
        "CR-506: Install a new cafeteria scheduling system in the hospital canteen to manage staff meal bookings and dietary preferences — this is a separate non-clinical system.",
    ])
    cr.save(os.path.join(OUTPUT_DIR, "06_Change_Requests_HospitalPMS.docx"))

    # ── 07 Meeting Minutes ──────────────────────────────────────────────────
    mom = create_doc("Meeting Minutes — Hospital Patient Management System", [
        "DEC-601: Team agreed that FR-101 patient registration must support both web browser and mobile app interfaces to improve patient onboarding rates.",
        "DEC-602: Security team decided that FR-105 credential handling must be reviewed by the compliance officer before FR-105 is implemented in any test environment.",
        "DEC-603: Architecture team decided that FR-109 audit logging must use a dedicated append-only PostgreSQL schema to prevent tampering and meet HIPAA audit requirements.",
        "DEC-604: Product team confirmed that FR-106 notification delivery must support international SMS via Twilio for patients registered with non-local phone numbers.",
        "DEC-605: Development team agreed to prioritize FR-102 appointment scheduling in Sprint 3 as it is a dependency for FR-106 reminder notifications.",
        "DEC-606: QA team agreed that TC-405 credential storage test must be executed on a staging environment with a sanitized patient dataset, not production data.",
        "DEC-607: Facilities team decided to procure 200 new ergonomic chairs and adjustable desks for the physician ward renovation — this is a physical facilities decision unrelated to the software system.",
    ])
    mom.save(os.path.join(OUTPUT_DIR, "07_Meeting_Minutes_HospitalPMS.docx"))

    print(f"[Hospital PMS] Generated 7 documents in: {OUTPUT_DIR}")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if f.endswith(".docx"):
            print(f"  ✓ {f}")


if __name__ == "__main__":
    main()
